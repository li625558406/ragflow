# test/test_flow_doc_table_edit.py
"""表格单元格编辑纯 helper 单测：不依赖 Quart/DB，直接构造 python-docx 文档。

说明：直接 `from api.apps.restful_apis.flow_app import ...` 会触发 api/apps/__init__.py
的 `settings.init_settings()`，进而需要本机 Redis/ES（非单元测试环境）。故照
test_permission_app.py 的模式：从源文件加载 flow_app 模块，并注入最小桩依赖
（仅 api.apps，避免 init_settings；其余模块均可真实导入）。
"""

import os
import sys
import types
from importlib.util import module_from_spec, spec_from_file_location

import pytest
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))


def _make_stub_module(name, **attrs):
    mod = types.ModuleType(name)
    for k, v in attrs.items():
        setattr(mod, k, v)
    sys.modules[name] = mod
    return mod


def _noop_decorator(*a, **kw):
    def deco(f):
        return f

    return deco


def _load_flow_app():
    _make_stub_module("api.apps", current_user=None, login_required=_noop_decorator)
    path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "api", "apps", "restful_apis", "flow_app.py"))
    spec = spec_from_file_location("flow_app_under_test", path)
    mod = module_from_spec(spec)
    sys.modules["flow_app_under_test"] = mod
    spec.loader.exec_module(mod)
    return mod


_flow_app = _load_flow_app()
_build_para_map = _flow_app._build_para_map
_apply_cell_text = _flow_app._apply_cell_text
_parse_table_edits = _flow_app._parse_table_edits


def _doc_with_table():
    doc = Document()
    doc.add_paragraph("前置段落")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).paragraphs[0].add_run("甲")
    t.cell(0, 1).paragraphs[0].add_run("乙")
    t.cell(1, 0).paragraphs[0].add_run("丙")
    t.cell(1, 1).paragraphs[0].add_run("丁")
    doc.add_paragraph("后置段落")
    return doc


def test_para_map_table_entry_is_docx_table():
    doc = _doc_with_table()
    pm = _build_para_map(doc)
    assert pm[0][0] == "p"
    assert pm[1][0] == "table"
    table = pm[1][1]
    assert table.cell(0, 0).text == "甲"
    assert pm[2][0] == "p"


def test_apply_cell_text_replace_and_clear():
    doc = _doc_with_table()
    table = _build_para_map(doc)[1][1]
    _apply_cell_text(table.cell(0, 0), "甲改", None)
    assert table.cell(0, 0).text == "甲改"
    # 清空：new_text 允许空串
    _apply_cell_text(table.cell(0, 1), "", None)
    assert table.cell(0, 1).text == ""


def test_apply_cell_text_runs_bold():
    doc = _doc_with_table()
    table = _build_para_map(doc)[1][1]
    runs = [{"text": "加粗", "bold": True}]
    _apply_cell_text(table.cell(1, 0), "加粗", runs)
    cell = table.cell(1, 0)
    assert cell.text == "加粗"
    assert cell.paragraphs[0].runs[0].bold is True


def test_apply_cell_text_multiline_writes_br_and_clears_extra_paras():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    cell.add_paragraph("第二段")
    assert len(cell.paragraphs) == 2  # 前置确认：单元格本就两段
    _apply_cell_text(cell, "一行\n两行", None)
    assert len(cell.paragraphs) == 2
    texts = [p.text for p in cell.paragraphs]
    assert texts[0].startswith("一行") and texts[1] == ""


def test_parse_table_edits_ok_and_empty_text_allowed():
    items = _parse_table_edits(
        [
            {"para_index": 1, "row": 0, "col": 0, "new_text": "改"},
            {"para_index": 1, "row": 0, "col": 1, "new_text": ""},
            {"para_index": 1, "row": 1, "col": 0, "new_text": "加粗", "runs": [{"text": "加粗", "bold": True}]},
        ]
    )
    assert len(items) == 3
    assert items[1]["new_text"] == ""


@pytest.mark.parametrize(
    "bad",
    [
        {"para_index": 1, "row": -1, "col": 0, "new_text": "x"},  # 负数
        {"para_index": 1, "row": 0, "col": "a", "new_text": "x"},  # 非整数
        {"para_index": 1, "row": 0, "col": 0, "new_text": "x" * 20001},  # 超长
        {"para_index": 1, "row": 0, "col": 0, "new_text": "x", "runs": [{"text": "y"}]},  # runs 不一致
        {"para_index": 1, "row": 0, "col": 0, "new_text": "", "runs": [{"text": "y"}]},  # 清空带 runs
        "not-a-dict",
    ],
)
def test_parse_table_edits_rejects_bad_input(bad):
    with pytest.raises(ValueError):
        _parse_table_edits([bad])
