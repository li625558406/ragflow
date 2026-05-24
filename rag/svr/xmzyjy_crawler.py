#!/usr/bin/env python3
#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
"""
Crawler for xmzyjy.cn (Xiamen Public Resources Trading Center).

Targets https://xmzyjy.cn/transactionInfo.html which is a Vue.js SPA with
REST API backend at /prod-api.  All listing and detail data is accessible
via direct API calls without Playwright.

Five modules, each with multiple sub-types:
  1. 工程建设 (13 sub-types)
  2. 政府采购 (6 sub-types)
  3. 行政事业资产资源交易 (6 sub-types)
  4. 土地矿业 (4 sub-types)
  5. 产权交易 (4 sub-types)

Date filtering: only items published today are collected.
Checkpoint/resume: each section is processed independently (list → details
→ upload → save state).  If the 3600s task timeout kills the run mid-way,
the next trigger resumes from the next incomplete section.

Usage:
    python xmzyjy_crawler.py \
        --tenant-id <TENANT_ID> \
        --target-url https://xmzyjy.cn/transactionInfo.html \
        --kb-id <KB_ID> \
        --task-name <NAME>
"""

import argparse
import json
import logging
import os
import random
import re
import sys
import time
import urllib.request
import urllib.error
import ssl
import zipfile
import io
from datetime import datetime
from urllib.parse import urljoin, urlparse, unquote
from bs4 import BeautifulSoup

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.abspath(os.path.join(_SCRIPT_DIR, "..", ".."))
sys.path.insert(0, _PROJECT_ROOT)

from common import settings
from common.log_utils import init_root_logger
from common.misc_utils import get_uuid

# ---------------------------------------------------------------------------
# Optional: Playwright for edge cases
# ---------------------------------------------------------------------------
try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
_API_BASE = "https://xmzyjy.cn/prod-api"
_SITE_ROOT = "https://xmzyjy.cn"
_FILE_SERVER = "https://xmzyjy.cn:7202/bidpublic_xm/synfile/"
_DOWNLOAD_BASE = "https://xmzyjy.cn:9100"

_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/120.0.0.0 Safari/537.36"
)

_REQUEST_DELAY_MIN = 1.0
_REQUEST_DELAY_MAX = 2.5

_STATE_FILENAME = "_crawler_state.json"

# SSL: allow self-signed / internal certs
_SSL_CONTEXT = ssl.create_default_context()
_SSL_CONTEXT.check_hostname = False
_SSL_CONTEXT.verify_mode = ssl.CERT_NONE

# Chrome paths for Playwright fallback
_CHROME_PATHS = [
    "C:/Program Files/Google/Chrome/Application/chrome.exe",
    "C:/Program Files (x86)/Google/Chrome/Application/chrome.exe",
    "/opt/chrome/chrome",
    "/usr/bin/google-chrome",
    "/usr/bin/chromium-browser",
    "/usr/bin/chromium",
]

# ---------------------------------------------------------------------------
# Module / sub-type definitions
# ---------------------------------------------------------------------------

PROJECT_SUBTYPES = [
    # (name, api_endpoint, date_field, detail_type, classPath/extra_params)
    ("\u62db\u6807\u8ba1\u5212", "tenderPlanList", "publishDate", "tenderPlan", None),
    ("\u8d44\u683c\u9884\u5ba1/\u62db\u6807\u516c\u544a", "quaInqueryAnnList", "noticeSendTime", "listing", None),
    ("\u53d8\u66f4\u516c\u544a", "qaList", "noticeSendTime", "listing", None),
    ("\u5f00\u6807\u8bb0\u5f55\u8868", "openingRecordList", "noticeSendTime", "listing", None),
    ("\u5b9a\u6807/\u4e2d\u6807\u5019\u9009\u4eba\u516c\u793a", "candidateAnnList", "noticeSendTime", "listing", None),
    ("\u4e2d\u6807\u7ed3\u679c\u516c\u793a", "winResultList", "noticeSendTime", "listing", None),
    ("\u901a\u77e5\u516c\u544a", "noticeList", "noticeSendTime", "listing", None),
    ("\u4fdd\u8bc1\u91d1\u9000\u8fd8", "sendBackMarginList", "sendBackTime", "listing", None),
    ("\u5f02\u8bae\u56de\u590d", "objectReplyList", "approveTime", "listing", None),
    ("\u903e\u671f\u4fdd\u8bc1\u91d1", "proOverdueBidbondList", "noticeSendTime", "listing", None),
    ("\u9879\u76ee\u96f7\u540c/\u7406\u8d54\u60c5\u51b5", "proIdenticalClaimList", "noticeSendTime", "listing", None),
    ("\u66dd\u5149\u4e13\u680f", "articleNewsList", "pubDate", "article", "p12"),
    ("\u653f\u7b56\u6cd5\u89c4", "articleNewsList", "pubDate", "article", "p3p1"),
]

PURCHASE_SUBTYPES = [
    ("\u91c7\u8d2d\u9879\u76ee\u516c\u544a", "quali_inquery_ann", None),
    ("\u91c7\u8d2d\u7ed3\u679c\u516c\u544a", "bid_deal_announce", None),
    ("\u91c7\u8d2d\u5408\u540c", "purchase_contract", None),
    ("\u66f4\u6b63\u516c\u544a", "correction_item", None),
    ("\u66dd\u5149\u4e13\u680f", "subjectBeehavior", None),
    ("\u653f\u7b56\u6cd5\u89c4", "articleNews_purchase", "p3p2"),
]

RECOURCE_SUBTYPES = [
    ("\u914d\u7f6e\u4fe1\u606f\u516c\u544a", "listing_pub_info", None),
    ("\u914d\u7f6e\u7ed3\u679c\u516c\u793a", "trade_result_info", None),
    ("\u66f4\u6b63\u516c\u544a", "other_pub_info", None),
    ("\u914d\u7f6e\u7a0b\u5e8f\u4fe1\u606f", "listing_pubcx_info", None),
    # ("\u7f51\u4e0a\u81ea\u7531\u7ade\u4ef7", "freeBidding", None),  # different system, skip
    ("\u653f\u7b56\u6cd5\u89c4", "articleNews_resource", "p3p3"),
]

LAND_SUBTYPES = [
    ("\u51fa\u8ba9\u516c\u544a", "dealBehaviorInfoList", "publishingTime", "article", None),
    ("\u4ea4\u6613\u7ed3\u679c", "landDealLandInfoList", "publishingTime", "article", None),
    ("\u53d8\u66f4\u516c\u544a", "qzhyzbycbgList", "approvalTime", "article", None),
    ("\u653f\u7b56\u6cd5\u89c4", "articleNewsList", "pubDate", "article", "p3p4"),
]

CQJY_SUBTYPES = [
    ("\u51fa\u8ba9\u516c\u544a", "HY811", None),
    ("\u4ea4\u6613\u7ed3\u679c", "HY812", None),
    ("\u53d8\u66f4\u516c\u544a", "HY814", None),
    ("\u7f51\u4e0a\u81ea\u7531\u7ade\u4ef7", "HY813", None),
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _safe_print(msg):
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode("gbk", errors="replace").decode("gbk"))


def _request_delay():
    time.sleep(random.uniform(_REQUEST_DELAY_MIN, _REQUEST_DELAY_MAX))


def _find_chrome():
    for p in _CHROME_PATHS:
        if os.path.exists(p):
            return p
    return None


def _today_str():
    return datetime.now().strftime("%Y-%m-%d")


def _date_match(d1, d2):
    """Compare two date strings ignoring time portion."""
    if not d1 or not d2:
        return False
    d1_clean = str(d1)[:10].replace("/", "-")
    d2_clean = str(d2)[:10].replace("/", "-")
    return d1_clean == d2_clean


# ---------------------------------------------------------------------------
# API helpers
# ---------------------------------------------------------------------------

def _api_headers(referer=None):
    return {
        "User-Agent": _USER_AGENT,
        "Content-Type": "application/json;charset=utf-8",
        "Accept": "application/json",
        "Referer": referer or "https://xmzyjy.cn/transactionInfo.html",
        "Origin": "https://xmzyjy.cn",
    }


def api_post(path, body, timeout=30):
    """POST JSON to /prod-api/path."""
    url = _API_BASE + path
    data = json.dumps(body, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers=_api_headers(), method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout, context=_SSL_CONTEXT) as resp:
            return json.loads(resp.read().decode("utf-8", errors="replace"))
    except urllib.error.HTTPError as e:
        logging.warning("API POST %s → HTTP %s", path, e.code)
        return {"code": e.code, "msg": str(e), "data": []}
    except Exception as e:
        logging.warning("API POST %s → %s", path, e)
        return {"code": -1, "msg": str(e), "data": []}


def api_get(path, params=None, timeout=30):
    """GET from /prod-api/path with optional query params."""
    url = _API_BASE + path
    if params:
        qs_parts = []
        for k, v in params.items():
            if v is not None:
                qs_parts.append("{}={}".format(k, urllib.request.quote(str(v))))
        if qs_parts:
            url += "?" + "&".join(qs_parts)
    req = urllib.request.Request(url, headers=_api_headers())
    try:
        with urllib.request.urlopen(req, timeout=timeout, context=_SSL_CONTEXT) as resp:
            return json.loads(resp.read().decode("utf-8", errors="replace"))
    except urllib.error.HTTPError as e:
        logging.warning("API GET %s → HTTP %s", path, e.code)
        return {"code": e.code, "msg": str(e), "data": []}
    except Exception as e:
        logging.warning("API GET %s → %s", path, e)
        return {"code": -1, "msg": str(e), "data": []}


def _download_file(url, timeout=60):
    """Download binary content from a URL. Returns (bytes, content_type) or (None, None)."""
    headers = {
        "User-Agent": _USER_AGENT,
        "Accept": "*/*",
        "Referer": "https://xmzyjy.cn/",
    }
    req = urllib.request.Request(url, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout, context=_SSL_CONTEXT) as resp:
            ct = resp.headers.get("Content-Type", "")
            return resp.read(), ct
    except Exception as e:
        logging.warning("Download failed for %s: %s", url, e)
        return None, None


# ---------------------------------------------------------------------------
# State management
# ---------------------------------------------------------------------------

def _load_state(output_dir):
    path = os.path.join(output_dir, _STATE_FILENAME)
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logging.warning("Failed to load crawler state: %s", e)
    return {"processed_ids": [], "completed_sections": []}


def _save_state(output_dir, state):
    path = os.path.join(output_dir, _STATE_FILENAME)
    os.makedirs(output_dir, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False)
    logging.info("State saved (%d IDs, %d sections done)",
                 len(state.get("processed_ids", [])),
                 len(state.get("completed_sections", [])))


# ---------------------------------------------------------------------------
# Markdown & attachment helpers
# ---------------------------------------------------------------------------

def _html_to_text(html_content):
    """Convert HTML content to plain text, preserving structure."""
    if not html_content:
        return ""
    try:
        soup = BeautifulSoup(html_content, "html.parser")
        # Remove script/style tags
        for tag in soup(["script", "style", "meta", "link"]):
            tag.decompose()
        # Convert br to newlines
        for br in soup.find_all("br"):
            br.replace_with("\n")
        # Convert p/div to double newlines
        for tag in soup.find_all(["p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"]):
            tag.append("\n")
        text = soup.get_text()
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines()]
        text = "\n".join(line for line in lines if line)
        return text
    except Exception as e:
        logging.warning("HTML to text conversion failed: %s", e)
        # Fallback: simple regex tag removal
        clean = re.sub(r'<[^>]+>', ' ', html_content)
        clean = re.sub(r'\s+', ' ', clean).strip()
        return clean


def _extract_attachment_urls(html_content):
    """Extract attachment/file URLs from HTML content."""
    urls = []
    if not html_content:
        return urls

    # Common patterns in the content
    for pattern in [
        r'href=["\']([^"\']*(?:\.doc|\.docx|\.pdf|\.zip|\.rar|\.xls|\.xlsx|\.jpg|\.png|\.gif|\.bmp|bidpublic_xm[^"\']*))["\']',
        r'src=["\']([^"\']*bidpublic_xm[^"\']*)["\']',
        r'["\'](https?://[^"\']*(?:download|file|attach|upload)[^"\']*)["\']',
    ]:
        for m in re.finditer(pattern, html_content, re.I):
            url = m.group(1)
            if url not in urls and not url.endswith(('.css', '.js')):
                urls.append(url)
    return urls


def _download_and_parse_attachment(url, output_dir):
    """Download an attachment and extract text content. Returns text string."""
    if not url.startswith("http"):
        if url.startswith("/"):
            url = _SITE_ROOT + url
        elif url.startswith("bidpublic_xm"):
            url = "https://xmzyjy.cn:7202/" + url
        else:
            url = _FILE_SERVER + url.lstrip("/")

    data, ct = _download_file(url)
    if not data:
        return ""

    fname = url.split("/")[-1].split("?")[0]
    ext = os.path.splitext(fname)[1].lower()

    if ext in ('.pdf',):
        return _parse_pdf(data, fname, output_dir)
    elif ext in ('.doc', '.docx'):
        return _parse_docx(data, fname, output_dir)
    elif ext in ('.zip', '.rar'):
        return _parse_zip(data, fname, output_dir)
    elif ext in ('.txt', '.md', '.csv', '.json', '.xml'):
        try:
            return data.decode('utf-8', errors='replace')[:50000]
        except Exception:
            return ""
    else:
        return ""


def _parse_pdf(data, fname, output_dir):
    """Extract text from PDF bytes using local OCR or pdfplumber."""
    # Try pdfplumber first
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages[:30]:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)[:100000]
    except ImportError:
        pass
    except Exception as e:
        logging.warning("pdfplumber failed for %s: %s", fname, e)

    # Try PyPDF2 / pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        text_parts = []
        for page in reader.pages[:30]:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        return "\n".join(text_parts)[:100000]
    except ImportError:
        pass
    except Exception as e:
        logging.warning("pypdf failed for %s: %s", fname, e)

    # OCR fallback
    try:
        from deepdoc.vision import OCR
        ocr = OCR()
        images = _pdf_to_images(data)
        if images:
            texts = []
            for img in images[:10]:
                t = ocr.ocr(img)
                if t:
                    texts.append(t)
            return "\n".join(texts)[:50000]
    except ImportError:
        pass
    except Exception as e:
        logging.warning("OCR fallback failed for %s: %s", fname, e)

    return ""


def _pdf_to_images(data):
    """Convert PDF bytes to PIL Images list."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        images = []
        for page in doc[:30]:
            pix = page.get_pixmap(dpi=200)
            from PIL import Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
        return images
    except ImportError:
        pass
    return []


def _parse_docx(data, fname, output_dir):
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(cells))
        return "\n".join(text_parts)[:100000]
    except ImportError:
        logging.warning("python-docx not installed, can't parse %s", fname)
    except Exception as e:
        logging.warning("DOCX parse failed for %s: %s", fname, e)
    return ""


def _parse_zip(data, fname, output_dir):
    """Extract and parse files from ZIP archive."""
    texts = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for name in zf.namelist():
                if name.endswith('/'):
                    continue
                try:
                    content = zf.read(name)
                    ext = os.path.splitext(name)[1].lower()
                    if ext in ('.pdf',):
                        t = _parse_pdf(content, name, output_dir)
                    elif ext in ('.docx', '.doc'):
                        t = _parse_docx(content, name, output_dir)
                    elif ext in ('.txt', '.md', '.csv', '.json', '.xml'):
                        t = content.decode('utf-8', errors='replace')[:50000]
                    else:
                        continue
                    if t:
                        texts.append("--- {} ---\n{}".format(name, t))
                except Exception as e:
                    logging.warning("Failed to extract %s from zip: %s", name, e)
        return "\n\n".join(texts)[:100000]
    except Exception as e:
        logging.warning("ZIP parse failed for %s: %s", fname, e)
    return ""


def _build_markdown(title, section, module_name, date_str, url, content_text, attachments_text):
    """Build a markdown document from extracted data."""
    parts = [
        "# {}".format(title or "\u65e0\u6807\u9898"),
        "",
        "**\u6a21\u5757:** {}".format(module_name),
        "**\u680f\u76ee:** {}".format(section),
        "**\u65e5\u671f:** {}".format(date_str or ""),
        "**URL:** {}".format(url or ""),
        "",
        "## \u6b63\u6587",
        "",
        content_text or "",
    ]

    if attachments_text:
        parts.extend([
            "",
            "## \u9644\u4ef6\u5185\u5bb9",
            "",
            attachments_text,
        ])

    parts.append("")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# KB upload
# ---------------------------------------------------------------------------

def _upload_to_kb(filepath, kb_id, tenant_id, parser_id="general"):
    from api.db.services.knowledgebase_service import KnowledgebaseService
    from api.db.services.file_service import FileService
    from api.db.services.document_service import DocumentService

    ok, kb = KnowledgebaseService.get_by_id(kb_id)
    if not ok:
        raise LookupError("Knowledge base {} not found".format(kb_id))

    with open(filepath, "rb") as f:
        blob = f.read()

    class _FileObj:
        def __init__(self, fn, b):
            self.id = get_uuid()
            self.filename = fn
            self.blob = b
        def read(self):
            return self.blob

    fo = _FileObj(os.path.basename(filepath), blob)
    errs, pairs = FileService.upload_document(kb, [fo], tenant_id)
    if errs:
        logging.warning("Upload errors: %s", errs)
    for doc, _ in pairs:
        did = doc["id"]
        try:
            DocumentService.update_by_id(did, {"parser_id": parser_id})
        except Exception:
            pass
        try:
            DocumentService.begin2parse(did)
            DocumentService.run(tenant_id, doc, {})
        except Exception as e:
            logging.error("Queue parse for %s: %s", did, e)


# ---------------------------------------------------------------------------
# Listing: paginate API and collect today's items
# ---------------------------------------------------------------------------

def _fetch_listing(api_path, params, date_field, today, max_pages=50):
    """Paginate through a listing API, return items matching today's date.

    Stops pagination early when records are all older than today (sorted desc by date).
    """
    items = []
    page = 1
    params = dict(params)
    params["pageSize"] = 20

    while page <= max_pages:
        params["pageNum"] = page
        r = api_post(api_path, params)
        if r.get("code") != 200:
            logging.warning("API %s page %d returned code %s", api_path, page, r.get("code"))
            break

        rows = r.get("data", [])
        if not rows:
            break

        found_today = False
        found_valid_older = False
        first_valid_date = None

        for row in rows:
            date_val = row.get(date_field, "")
            if not date_val or str(date_val).strip() in ("", "1", "None"):
                continue

            if _date_match(date_val, today):
                found_today = True
                items.append(row)
            else:
                # Track the most recent valid date
                d10 = str(date_val)[:10]
                if d10 < today and d10[:4].isdigit():
                    found_valid_older = True
                if first_valid_date is None:
                    first_valid_date = d10

        # Stop if: (a) page has older items and no today items, or
        # (b) page 1's first valid date is before today (no today items exist at all)
        if found_valid_older and not found_today:
            break
        if page == 1 and first_valid_date and first_valid_date < today:
            # First page oldest date is before today → no today items exist
            break
        if page == 1 and not found_today and not first_valid_date:
            # No valid dates found on first page → assume no today items
            break

        total = r.get("totalNum", 0)
        if page * params["pageSize"] >= total:
            break

        page += 1
        _request_delay()

    return items


# ---------------------------------------------------------------------------
# Detail fetching
# ---------------------------------------------------------------------------

def _item_to_display_name(item):
    """Get a human-readable name/title from a listing item."""
    return str(item.get("tenderProjectName")
           or item.get("projname")
           or item.get("projectName")
           or item.get("announcementTitle")
           or item.get("title")
           or item.get("id", ""))


def _item_to_content_text(item):
    """Convert a listing row to formatted text content."""
    lines = []
    for k, v in item.items():
        if k.startswith("_") or v is None:
            continue
        if k in ("id", "guid", "pcodeTprojectUniqueId", "objId", "dataKey",
                 "uplink", "blockHeight", "witnessId", "transactionNo",
                 "subgroupid", "isDeleted", "isRevoke", "jyId", "zyjyKeyId"):
            continue
        # Format the value
        if isinstance(v, (int, float)):
            lines.append("{}: {}".format(k, v))
        elif isinstance(v, str) and len(v) > 200:
            lines.append("{}: {}...".format(k, v[:200]))
        elif isinstance(v, list):
            lines.append("{}: {}".format(k, ", ".join(str(x) for x in v[:10])))
        else:
            lines.append("{}: {}".format(k, v))
    return "\n".join(lines)


def _fetch_detail_article(item_id, class_path=None):
    """Fetch article detail via GET /api/news/info (for article/news types only)."""
    params = {"id": str(item_id)}
    if class_path:
        params["classPath"] = class_path
    r = api_get("/api/news/info", params)
    if r.get("code") == 200 and r.get("data"):
        data = r["data"]
        if isinstance(data, dict):
            inner = data.get("data", data)
            return {
                "title": inner.get("title", ""),
                "content_text": _html_to_text(inner.get("content", "")),
                "date_str": inner.get("createTime", inner.get("pubDate", "")),
            }
    return {"title": "", "content_text": "", "date_str": ""}


def _fetch_detail_tender_plan(item_id):
    """Fetch tender plan detail via specific API."""
    r = api_get("/api/trade/project/getProTenderPlanDetail", {"id": str(item_id)})
    if r.get("code") == 200 and r.get("data"):
        inner = r["data"][0] if isinstance(r["data"], list) and r["data"] else r["data"]
        if isinstance(inner, dict):
            return {
                "title": inner.get("tenderProjectName", ""),
                "content_text": _item_to_content_text(inner),
                "date_str": inner.get("publishDate", ""),
                "content_html": json.dumps(inner, ensure_ascii=False),
            }
    return {"title": "", "content_text": "", "date_str": "", "content_html": ""}


def _fetch_detail_from_item(item, detail_type, class_path=None):
    """Fetch detail for an item based on its detail type.

    - "tenderPlan": use getProTenderPlanDetail API
    - "article": use /api/news/info
    - "listing": use the item data itself as content
    """
    item_id = str(item.get("id", ""))

    if detail_type == "tenderPlan":
        return _fetch_detail_tender_plan(item_id)
    elif detail_type == "article":
        return _fetch_detail_article(item_id, class_path)
    else:
        # Use listing data as content
        name = _item_to_display_name(item)
        date_str = item.get("publishDate") or item.get("pubDate") or item.get("publishtime") or item.get("publishingTime") or item.get("noticeSendTime") or item.get("sendBackTime") or item.get("approveTime") or ""
        return {
            "title": name,
            "content_text": _item_to_content_text(item),
            "date_str": str(date_str)[:10],
        }


# ---------------------------------------------------------------------------
# Section-level processing
# ---------------------------------------------------------------------------

def _process_items(items, section_key, module_name, detail_type, class_path,
                   output_dir, kb_id, tenant_id, processed_ids, state,
                   start_time, max_runtime):
    """Process a list of items: fetch detail, build markdown, upload, checkpoint."""
    new_items = []
    for item in items:
        item_id = "{}|{}".format(section_key, item.get("id", ""))
        if item_id not in processed_ids:
            new_items.append(item)

    if not new_items:
        _safe_print("[{}] 0 new items.".format(section_key))
        return 0

    _safe_print("[{}] {} new items, processing...".format(section_key, len(new_items)))
    sys.stdout.flush()

    BATCH_SIZE = 10
    total_processed = 0
    batch_num = 0

    for batch_start in range(0, len(new_items), BATCH_SIZE):
        batch = new_items[batch_start:batch_start + BATCH_SIZE]
        batch_num += 1
        md_parts = []
        batch_ids = []

        for idx, item in enumerate(batch):
            # Check timeout
            elapsed = time.time() - start_time
            remaining = max_runtime - elapsed
            if remaining < 120:
                _safe_print("[{}] Runtime {:.0f}s, stopping early. Saved {} records.".format(
                    section_key, elapsed, total_processed))
                sys.stdout.flush()
                processed_ids.update(batch_ids)
                state["processed_ids"] = list(processed_ids)
                _save_state(output_dir, state)
                return total_processed

            global_idx = batch_start + idx + 1
            item_id_raw = str(item.get("id", ""))
            full_id = "{}|{}".format(section_key, item_id_raw)

            row_name = _item_to_display_name(item)[:60]
            _safe_print("[{}] [{}/{}] {} (id={})".format(
                section_key, global_idx, len(new_items), row_name, item_id_raw))
            sys.stdout.flush()

            # Fetch detail
            detail = _fetch_detail_from_item(item, detail_type, class_path)

            content_text = detail.get("content_text", "")
            if not content_text:
                content_text = _item_to_content_text(item)

            # Extract attachments from content_html if available
            attachment_urls = []
            html = detail.get("content_html", "")
            if html:
                attachment_urls = _extract_attachment_urls(html)

            # Also check downloadUrl field
            dl_url = item.get("downloadUrl", "")
            if dl_url and dl_url.startswith("http") and "9103" not in dl_url:
                attachment_urls.append(dl_url)

            attachment_texts = []
            for att_url in attachment_urls[:5]:
                _safe_print("   Downloading attachment: {}".format(att_url[:80]))
                att_text = _download_and_parse_attachment(att_url, output_dir)
                if att_text:
                    attachment_texts.append(att_text)
                _request_delay()

            date_str = detail.get("date_str", "")
            if not date_str:
                date_str = str(item.get("publishDate") or item.get("pubDate")
                            or item.get("publishtime") or item.get("publishingTime")
                            or item.get("noticeSendTime") or item.get("sendBackTime")
                            or item.get("approveTime") or "")[:10]

            # Build markdown
            md = _build_markdown(
                title=detail.get("title", row_name),
                section=section_key,
                module_name=module_name,
                date_str=date_str,
                url="https://xmzyjy.cn/noticeDetail.html?id={}".format(item_id_raw),
                content_text=content_text,
                attachments_text="\n\n".join(attachment_texts) if attachment_texts else "",
            )

            # Save individual markdown
            safe_id = full_id.replace("|", "_").replace("/", "_")
            md_path = os.path.join(output_dir, "articles", "{}.md".format(safe_id))
            os.makedirs(os.path.dirname(md_path), exist_ok=True)
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md)

            md_parts.append(md)
            batch_ids.append(full_id)

            _request_delay()

        # Batch checkpoint
        if md_parts:
            safe_sec = section_key.replace("|", "_").replace("/", "_")
            batch_path = os.path.join(output_dir, "{}_{:03d}.md".format(safe_sec, batch_num))
            with open(batch_path, "w", encoding="utf-8") as f:
                f.write("\n\n---\n\n".join(md_parts))

            processed_ids.update(batch_ids)
            state["processed_ids"] = list(processed_ids)
            _save_state(output_dir, state)

            if kb_id:
                try:
                    _upload_to_kb(batch_path, kb_id, tenant_id)
                except Exception as e:
                    _safe_print("[{}] batch {} upload failed: {}".format(section_key, batch_num, e))
                    logging.error("Upload failed for %s batch %d: %s", section_key, batch_num, e)

            total_processed += len(md_parts)
            _safe_print("[{}] batch {} done ({}/{} total)".format(
                section_key, batch_num, total_processed, len(new_items)))
            sys.stdout.flush()

    return total_processed


# ---------------------------------------------------------------------------
# Module crawlers
# ---------------------------------------------------------------------------

def _crawl_project(output_dir, kb_id, tenant_id, processed_ids, state,
                   start_time, max_runtime, today):
    """Crawl 工程建设 (13 sub-types)."""
    module_name = "\u5de5\u7a0b\u5efa\u8bbe"
    _safe_print("\n=== {} ===\n".format(module_name))
    sys.stdout.flush()

    for st in PROJECT_SUBTYPES:
        name, endpoint, date_field, detail_type, class_path = st
        section_key = "gcjs|{}".format(name)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key in state.get("completed_sections", []):
            _safe_print("[{}] SKIPPED (already completed)".format(section_key))
            continue

        _safe_print("\n--- {} ---".format(name))

        if endpoint == "articleNewsList":
            params = {"pageSize": 20, "classPath": class_path}
            api_path = "/api/news/articleNewsList"
            items = _fetch_listing(api_path, params, date_field, today)
        else:
            params = {"pageSize": 20, "isXmStrack": ""}
            api_path = "/api/trade/project/{}".format(endpoint)
            items = _fetch_listing(api_path, params, date_field, today)

        if not items:
            _safe_print("[{}] 0 items today, marking done.".format(section_key))
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)
            continue

        # Skip revoked tender plans
        if endpoint == "tenderPlanList":
            items = [it for it in items if str(it.get("isRevoke", "0")) != "1"]

        _process_items(items, section_key, module_name, detail_type, class_path,
                      output_dir, kb_id, tenant_id, processed_ids, state,
                      start_time, max_runtime)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key not in state.get("completed_sections", []):
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)


def _crawl_purchase(output_dir, kb_id, tenant_id, processed_ids, state,
                    start_time, max_runtime, today):
    """Crawl 政府采购 (6 sub-types)."""
    module_name = "\u653f\u5e9c\u91c7\u8d2d"
    _safe_print("\n=== {} ===\n".format(module_name))
    sys.stdout.flush()

    for name, data_type, class_path in PURCHASE_SUBTYPES:
        section_key = "zfcg|{}".format(name)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key in state.get("completed_sections", []):
            _safe_print("[{}] SKIPPED (already completed)".format(section_key))
            continue

        _safe_print("\n--- {} ---".format(name))

        if data_type == "subjectBeehavior":
            api_path = "/api/trade/purchase/subjectBeehaviorList"
            params = {"pageSize": 20}
            items = _fetch_listing(api_path, params, "date", today)
        elif data_type.startswith("articleNews"):
            api_path = "/api/news/articleNewsList"
            params = {"pageSize": 20, "classPath": class_path}
            items = _fetch_listing(api_path, params, "pubDate", today)
        else:
            api_path = "/api/trade/purchase/purchaseList"
            params = {"pageSize": 20, "dataType": data_type}
            items = _fetch_listing(api_path, params, "publishtime", today)

        if not items:
            _safe_print("[{}] 0 items today, marking done.".format(section_key))
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)
            continue

        # Use listing data for all purchase types
        detail_type = "article" if data_type.startswith("articleNews") else "listing"
        _process_items(items, section_key, module_name, detail_type, class_path,
                      output_dir, kb_id, tenant_id, processed_ids, state,
                      start_time, max_runtime)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key not in state.get("completed_sections", []):
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)


def _crawl_recource(output_dir, kb_id, tenant_id, processed_ids, state,
                    start_time, max_runtime, today):
    """Crawl 行政事业资产资源交易 (6 sub-types)."""
    module_name = "\u884c\u653f\u4e8b\u4e1a\u8d44\u4ea7\u8d44\u6e90\u4ea4\u6613"
    _safe_print("\n=== {} ===\n".format(module_name))
    sys.stdout.flush()

    for name, data_type, class_path in RECOURCE_SUBTYPES:
        section_key = "xzsyzc|{}".format(name)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key in state.get("completed_sections", []):
            _safe_print("[{}] SKIPPED (already completed)".format(section_key))
            continue

        _safe_print("\n--- {} ---".format(name))

        if data_type.startswith("articleNews"):
            api_path = "/api/news/articleNewsList"
            params = {"pageSize": 20, "classPath": class_path}
            items = _fetch_listing(api_path, params, "pubDate", today)
        else:
            api_path = "/api/trade/recource/recourceConfigList"
            params = {"pageSize": 20, "dataType": data_type}
            items = _fetch_listing(api_path, params, "publishTime", today)

        if not items:
            _safe_print("[{}] 0 items today, marking done.".format(section_key))
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)
            continue

        detail_type = "article" if data_type.startswith("articleNews") else "listing"
        _process_items(items, section_key, module_name, detail_type, class_path,
                      output_dir, kb_id, tenant_id, processed_ids, state,
                      start_time, max_runtime)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key not in state.get("completed_sections", []):
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)


def _crawl_land(output_dir, kb_id, tenant_id, processed_ids, state,
                start_time, max_runtime, today):
    """Crawl 土地矿业 (4 sub-types)."""
    module_name = "\u571f\u5730\u77ff\u4e1a"
    _safe_print("\n=== {} ===\n".format(module_name))
    sys.stdout.flush()

    for name, endpoint, date_field, detail_type, class_path in LAND_SUBTYPES:
        section_key = "tdky|{}".format(name)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key in state.get("completed_sections", []):
            _safe_print("[{}] SKIPPED (already completed)".format(section_key))
            continue

        _safe_print("\n--- {} ---".format(name))

        if endpoint == "articleNewsList":
            api_path = "/api/news/articleNewsList"
            params = {"pageSize": 20, "classPath": class_path}
            items = _fetch_listing(api_path, params, date_field, today)
        else:
            api_path = "/api/trade/land/{}".format(endpoint)
            params = {"pageSize": 20}
            items = _fetch_listing(api_path, params, date_field, today)

        if not items:
            _safe_print("[{}] 0 items today, marking done.".format(section_key))
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)
            continue

        _process_items(items, section_key, module_name, detail_type, class_path,
                      output_dir, kb_id, tenant_id, processed_ids, state,
                      start_time, max_runtime)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key not in state.get("completed_sections", []):
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)


def _crawl_cqjy(output_dir, kb_id, tenant_id, processed_ids, state,
                start_time, max_runtime, today):
    """Crawl 产权交易 (4 sub-types)."""
    module_name = "\u4ea7\u6743\u4ea4\u6613"
    _safe_print("\n=== {} ===\n".format(module_name))
    sys.stdout.flush()

    for name, data_no, class_path in CQJY_SUBTYPES:
        section_key = "cqjy|{}".format(name)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key in state.get("completed_sections", []):
            _safe_print("[{}] SKIPPED (already completed)".format(section_key))
            continue

        _safe_print("\n--- {} ---".format(name))

        api_path = "/api/cqjy/pro/cqjyProInfo"
        params = {"pageSize": 20, "dataNo": data_no}
        items = _fetch_listing(api_path, params, "pubDate", today)

        if not items:
            _safe_print("[{}] 0 items today, marking done.".format(section_key))
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)
            continue

        _process_items(items, section_key, module_name, "listing", class_path,
                      output_dir, kb_id, tenant_id, processed_ids, state,
                      start_time, max_runtime)

        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            return

        if section_key not in state.get("completed_sections", []):
            state.setdefault("completed_sections", []).append(section_key)
            _save_state(output_dir, state)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args():
    p = argparse.ArgumentParser(description="xmzyjy.cn crawler")
    p.add_argument("--tenant-id", required=True)
    p.add_argument("--target-url", default="https://xmzyjy.cn/transactionInfo.html")
    p.add_argument("--kb-id", required=True)
    p.add_argument("--task-name", required=True)
    p.add_argument("--output-dir", default=None)
    p.add_argument("--full", action="store_true", help="Ignore state, re-crawl all")
    p.add_argument("--section", default=None,
                   help="Comma-separated: gcjs,zfcg,xzsyzc,tdky,cqjy (default: all)")
    p.add_argument("--max-runtime", type=int, default=3300,
                   help="Max runtime in seconds (default: 3300)")
    for opt in ("--llm-id", "--llm-model"):
        p.add_argument(opt, default=None)
    return p.parse_args()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    args = parse_args()

    _safe_print("\n" + "=" * 60)
    _safe_print("[XMZYJY] xmzyjy.cn crawler")
    _safe_print("[XMZYJY] KB: {}".format(args.kb_id))
    _safe_print("[XMZYJY] Max runtime: {}s".format(args.max_runtime))
    _safe_print("=" * 60 + "\n")
    sys.stdout.flush()

    settings.init_settings()
    logging.info("=== XMZYJY crawler started ===")

    output_dir = args.output_dir or os.path.join(
        _PROJECT_ROOT, "rag", args.task_name.strip()
    )
    os.makedirs(output_dir, exist_ok=True)
    _safe_print("[XMZYJY] Output: {}\n".format(output_dir))
    sys.stdout.flush()

    # Load state
    state = _load_state(output_dir) if not args.full else {
        "processed_ids": [], "completed_sections": []
    }
    processed_ids = set(state.get("processed_ids", []))
    _safe_print("[XMZYJY] Previously processed: {} IDs, completed sections: {}\n".format(
        len(processed_ids), len(state.get("completed_sections", []))))
    sys.stdout.flush()

    today = _today_str()
    _safe_print("[XMZYJY] Today: {}\n".format(today))
    sys.stdout.flush()

    start_time = time.time()
    max_runtime = args.max_runtime

    # Determine which modules to crawl
    sections = args.section.split(",") if args.section else ["gcjs", "zfcg", "xzsyzc", "tdky", "cqjy"]
    sections = [s.strip() for s in sections]

    for sec in sections:
        elapsed = time.time() - start_time
        if max_runtime - elapsed < 120:
            _safe_print("\n[XMZYJY] Runtime {:.0f}s, stopping. Remaining sections saved for next run.\n".format(elapsed))
            break

        if sec == "gcjs":
            _crawl_project(output_dir, args.kb_id, args.tenant_id,
                          processed_ids, state, start_time, max_runtime, today)
        elif sec == "zfcg":
            _crawl_purchase(output_dir, args.kb_id, args.tenant_id,
                           processed_ids, state, start_time, max_runtime, today)
        elif sec == "xzsyzc":
            _crawl_recource(output_dir, args.kb_id, args.tenant_id,
                           processed_ids, state, start_time, max_runtime, today)
        elif sec == "tdky":
            _crawl_land(output_dir, args.kb_id, args.tenant_id,
                       processed_ids, state, start_time, max_runtime, today)
        elif sec == "cqjy":
            _crawl_cqjy(output_dir, args.kb_id, args.tenant_id,
                       processed_ids, state, start_time, max_runtime, today)
        else:
            _safe_print("[XMZYJY] Unknown section: {}".format(sec))

    _safe_print("\n" + "=" * 60)
    _safe_print("[XMZYJY] Done. Total unique IDs: {}".format(len(processed_ids)))
    _safe_print("=" * 60 + "\n")
    sys.stdout.flush()
    logging.info("=== XMZYJY crawler finished ===")


if __name__ == "__main__":
    CONSUMER_NAME = "xmzyjy_crawler"
    init_root_logger(CONSUMER_NAME)
    main()
