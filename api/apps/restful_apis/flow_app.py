#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
"""C端流程（文件流转工作流）REST API

路由前缀：/api/v1/flow

端点：
  - POST   /flow                                    创建流程（发起人上传 v1 文件）
  - GET    /flow/list?scope=                        流程列表（todo/initiated/joined/all）
  - GET    /flow/<flow_id>                          流程详情（版本/批注/AI记录/视角）
  - POST   /flow/<flow_id>/version                  追加新版本（人工上传）
  - GET    /flow/<flow_id>/version/<version_id>/download  下载某版本文件
  - POST   /flow/<flow_id>/comment                  添加批注意见（通知其他参与人）
  - POST   /flow/<flow_id>/ai-record                记录一次 AI 处理（回复可落为新版本）
  - POST   /flow/<flow_id>/document/edit            编辑文档段落（存为新版本，仅当前节点负责人）
  - POST   /flow/<flow_id>/submit                   流转（next）/ 退回（return）
  - POST   /flow/<flow_id>/archive                  归档（仅汇总节点发起人）
  - POST   /flow/<flow_id>/cancel                   作废（仅发起人）
"""
import hashlib
import logging
import os
import re
import time
import uuid
from io import BytesIO
from urllib.parse import quote

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph
from quart import Blueprint, Response, request

from api.apps import current_user, login_required
from api.db.db_models import FlowInstance, User
from api.db.services.flow_service import (
    FlowActionService,
    FlowAiChatService,
    FlowCommentService,
    FlowInstanceService,
    FlowVersionService,
    FlowWorkflow,
    _bucket_of,
    notify_flow_event,
    notify_target_of,
)
from api.utils.api_utils import get_json_result
from api.utils.doc_utils import doc_to_docx_via_libreoffice, is_doc_file
from common import settings
from common.misc_utils import thread_pool_exec

manager = Blueprint("rest_flow_app", __name__)

logger = logging.getLogger(__name__)

_SCOPES = ("todo", "initiated", "joined", "all")


def _err(msg: str, code: int = 100):
    return get_json_result(code=code, message=msg)


def _flow_dict(flow_id: str):
    return FlowInstanceService.get_flow(flow_id)


def _require_participant(flow) -> dict:
    if not flow:
        raise LookupError("流程不存在")
    if not FlowWorkflow.can_view(flow, current_user.id):
        raise PermissionError("无权访问该流程")
    return flow


def _require_owner(flow) -> dict:
    _require_participant(flow)
    if FlowWorkflow.owner_of_current(flow) != current_user.id:
        raise PermissionError("当前不在你的节点上，无法操作")
    return flow


def _safe_filename(name: str) -> str:
    """剥掉路径分隔符与不可打印控制字符，避免 object name 被前端文件名注入子目录；超长时保留扩展名截断。"""
    base = os.path.basename((name or "").replace("\\", "/"))
    cleaned = "".join(ch for ch in base if ch.isprintable()).strip()
    if not cleaned:
        return "unnamed"
    if len(cleaned) > 200:
        root, ext = os.path.splitext(cleaned)
        digest = hashlib.md5(cleaned.encode("utf-8")).hexdigest()[:8]
        cleaned = f"{root[:160]}_{digest}{ext}"
    return cleaned


def _others_of(flow: dict, me: str) -> list:
    return [uid for uid in (flow["initiator_id"], flow["leader_id"], flow["handler_id"]) if uid != me]


def _nickname_of(uid: str) -> str:
    """通知文案用昵称展示，查不到时退回原始 id。"""
    try:
        u = User.get_or_none(User.id == uid)
        return (u.nickname or uid) if u else uid
    except Exception:
        return uid


def _action_error(e: Exception):
    """业务异常分级：PermissionError→403 / RuntimeError(乐观锁冲突)→409 / ValueError→100。"""
    if isinstance(e, PermissionError):
        return _err(str(e), 403)
    if isinstance(e, RuntimeError):
        return _err(str(e), 409)
    if isinstance(e, ValueError):
        return _err(str(e))
    return None


# ── 1. 创建流程 ────────────────────────────────────────────────────
@manager.route("/flow", methods=["POST"])  # noqa: F821
@login_required
async def create_flow():
    try:
        form = await request.form
        files = await request.files
        title = (form.get("title") or "").strip()
        leader_id = (form.get("leader_id") or "").strip()
        handler_id = (form.get("handler_id") or "").strip()
        file = files.get("file")
        if not title:
            return _err("标题不能为空", 101)
        if not leader_id:
            return _err("请选择领导", 101)
        if not handler_id:
            return _err("请选择处理人", 101)
        # 初始文件可选：不传则创建无版本的流程，由后续节点补传
        blob = None
        file_name = ""
        if file is not None and file.filename:
            # Quart FileStorage.read() 是同步方法（返回 bytes），放线程池避免阻塞事件循环
            blob = await thread_pool_exec(file.read)
            if not blob:
                return _err("文件内容为空", 101)
            file_name = _safe_filename(file.filename)
        if leader_id == current_user.id or handler_id == current_user.id:
            return _err("领导和处理人不能是发起人自己", 101)
        if leader_id == handler_id:
            return _err("领导和处理人不能是同一个人", 101)
        for uid, label in ((leader_id, "领导"), (handler_id, "处理人")):
            u = User.get_or_none(User.id == uid)
            if not u or u.status != "1":
                return _err(f"所选{label}不存在或已停用", 101)

        flow = FlowInstanceService.insert(
            title=title,
            initiator_id=current_user.id,
            leader_id=leader_id,
            handler_id=handler_id,
            status="initiator",
            current_version_id="",
        )
        flow_id = flow.id

        if blob is None:
            return get_json_result(data={"id": flow_id, "version": None})

        object_name = f"flow/{flow_id}/v1_{file_name}"
        try:
            await thread_pool_exec(settings.STORAGE_IMPL.put, _bucket_of(flow.__data__), object_name, blob)
        except Exception as e:
            logger.exception("flow storage put failed, rollback flow row %s", flow_id)
            FlowInstance.delete().where(FlowInstance.id == flow_id).execute()
            return _err(f"文件存储失败: {e}")
        flow_dict = _flow_dict(flow_id)
        try:
            version = FlowVersionService.add_version(
                flow_dict, object_name, file_name, file.mimetype or "", len(blob),
                "manual_upload", current_user.id,
            )
        except Exception as e:
            logger.exception("flow add_version failed, rollback flow row %s", flow_id)
            FlowInstance.delete().where(FlowInstance.id == flow_id).execute()
            return _err(f"版本记录失败: {e}")
        return get_json_result(data={"id": flow_id, "version": version})
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 2. 流程列表 ────────────────────────────────────────────────────
@manager.route("/flow/list", methods=["GET"])  # noqa: F821
@login_required
async def list_flows():
    try:
        scope = request.args.get("scope", "all")
        if scope not in _SCOPES:
            return _err(f"非法 scope: {scope}，可选值 todo/initiated/joined/all", 101)
        items, total = FlowInstanceService.list_for_user(current_user.id, scope)
        return get_json_result(data={"list": items, "total": total})
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 3. 流程详情 ────────────────────────────────────────────────────
@manager.route("/flow/<flow_id>", methods=["GET"])  # noqa: F821
@login_required
async def get_flow(flow_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        return get_json_result(data={
            "flow": flow,
            "versions": FlowVersionService.list_by_flow(flow_id),
            "comments": FlowCommentService.list_by_flow(flow_id),
            "ai_chats": FlowAiChatService.list_by_flow(flow_id),
            "viewer": {
                "is_owner": FlowWorkflow.owner_of_current(flow) == current_user.id,
                "is_initiator": flow["initiator_id"] == current_user.id,
            },
        })
    except LookupError as e:
        return _err(str(e), 404)
    except PermissionError as e:
        return _err(str(e), 403)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 3.1 流程参与人候选（所有启用用户，仅登录即可；发起流程选领导/处理人用） ──
@manager.route("/flow/candidates", methods=["GET"])  # noqa: F821
@login_required
async def flow_candidates():
    try:
        users = (
            User.select(User.id, User.nickname, User.email)
            .where(User.status == "1")
            .order_by(User.create_time)
        )
        return get_json_result(
            data={
                "list": [
                    {"id": u.id, "nickname": u.nickname or u.id}
                    for u in users
                ]
            }
        )
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 4. 追加新版本 ─────────────────────────────────────────────────
@manager.route("/flow/<flow_id>/version", methods=["POST"])  # noqa: F821
@login_required
async def upload_version(flow_id: str):
    try:
        flow = _require_owner(_flow_dict(flow_id))
        if flow["status"] in FlowWorkflow.TERMINAL:
            return _err("流程已结束")
        files = await request.files
        file = files.get("file")
        if file is None or not file.filename:
            return _err("请上传文件", 101)
        # 同 create_flow：FileStorage.read() 同步，放线程池
        blob = await thread_pool_exec(file.read)
        if not blob:
            return _err("文件内容为空", 101)

        file_name = _safe_filename(file.filename)
        # 外层先算一次 version_no 仅用于拼 object_name；add_version 内部会再算一次，
        # 无并发时两次结果一致（唯一索引 flow_id+version_no 兜底极端并发）。
        no = FlowVersionService.next_version_no(flow_id)
        object_name = f"flow/{flow_id}/v{no}_{file_name}"
        await thread_pool_exec(settings.STORAGE_IMPL.put, _bucket_of(flow), object_name, blob)
        version = FlowVersionService.add_version(
            flow, object_name, file_name, file.mimetype or "", len(blob),
            "manual_upload", current_user.id,
        )
        return get_json_result(data={"version": version})
    except LookupError as e:
        return _err(str(e), 404)
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 4.1 编辑文档段落（存为新版本） ────────────────────────────────
_CTRL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _build_para_map(doc):
    """复刻 rag/app/naive.py Docx.to_paragraphs 的遍历规则，建立
    para_index → ('p', DocxParagraph) / ('table', None) / ('image', None) 映射。
    w:p 空文本且无图跳过（不占 index）、有图记 image 占 index；w:tbl 整表占一个 index。"""
    para_map = {}
    idx = 0
    for block in doc.element.body:
        if block.tag.endswith("p"):
            p = DocxParagraph(block, doc)
            text = _CTRL_CHARS.sub("", (p.text or "").strip())
            if text:
                para_map[idx] = ("p", p)
                idx += 1
            else:
                # 空文本段落：与 naive.get_picture 一致，检测 pic:pic + a:blip 可解析关系才算图片；
                # 空文本且无图则完全跳过（不占 index，与 naive.to_paragraphs 的 continue 一致）
                has_img = False
                for img in block.xpath(".//pic:pic"):
                    embed = img.xpath(".//a:blip/@r:embed")
                    if embed and embed[0] in doc.part.related_parts:
                        has_img = True
                        break
                if has_img:
                    para_map[idx] = ("image", None)
                    idx += 1
        elif block.tag.endswith("tbl"):
            para_map[idx] = ("table", None)
            idx += 1
    return para_map


_COLOR_RE = re.compile(r"^#[0-9A-Fa-f]{6}$")
_RUN_BOOL_KEYS = ("bold", "italic", "underline", "strike", "superscript", "subscript")


def _parse_runs(raw):
    """解析并校验可选 runs 字段：None → None（走旧整段替换）；
    非法结构/颜色/字号抛 ValueError，由调用方转 400。"""
    if raw is None:
        return None
    if not isinstance(raw, list) or not raw:
        raise ValueError("runs 必须是非空数组或省略")
    if len(raw) > 500:
        raise ValueError("runs 片段数量超限（最多 500）")
    parsed = []
    for r in raw:
        if not isinstance(r, dict):
            raise ValueError("runs 项格式非法")
        text = _CTRL_CHARS.sub("", str(r.get("text") or ""))
        if not text:
            raise ValueError("runs 片段文本不能为空")
        item = {"text": text}
        for k in _RUN_BOOL_KEYS:
            if r.get(k):
                item[k] = True
        for k in ("color", "bg_color"):
            v = r.get(k)
            if v:
                v = str(v)
                if not _COLOR_RE.match(v):
                    raise ValueError(f"{k} 颜色值非法：{v}")
                item[k] = v
        font = r.get("font")
        if font:
            item["font"] = str(font)[:50]
        size = r.get("size")
        if size is not None:
            try:
                size = float(size)
            except (TypeError, ValueError):
                raise ValueError("size 字号必须是数字")
            if not (1 <= size <= 200):
                raise ValueError("size 字号超出范围（1-200pt）")
            item["size"] = size
        parsed.append(item)
    return parsed


def _replace_para_text(p: DocxParagraph, new_text: str):
    """整段替换文本：保留首 run 格式（沿用段落样式），其余 run 清空。
    超链接/域先移除——p.runs 不覆盖其内部 run，残留会导致新旧文本拼接。"""
    el = p._element
    for child in list(el):
        if child.tag in (qn("w:hyperlink"), qn("w:fldSimple")):
            el.remove(child)
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def _set_run_font(run, name: str):
    """同时设置西文（ascii/hAnsi）与中文（eastAsia）字体，否则中文不生效。"""
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def _apply_runs(p: DocxParagraph, runs):
    """按 runs 重写段落文本 run（保留段落级 style/对齐）。
    runs 经 _parse_runs 校验。bg_color 用 w:shd 底纹实现。"""
    # 清空段落内联内容：除直接 run 外还要移除超链接/域（p.runs 不覆盖它们，
    # 残留会拼接进重写后的段落）；保留 pPr 段落属性
    el = p._element
    for child in list(el):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:fldSimple")):
            el.remove(child)
    for item in runs:
        run = p.add_run(item["text"])
        if item.get("bold"):
            run.bold = True
        if item.get("italic"):
            run.italic = True
        if item.get("underline"):
            run.underline = True
        if item.get("strike"):
            run.font.strike = True
        if item.get("superscript"):
            run.font.superscript = True
        if item.get("subscript"):
            run.font.subscript = True
        if item.get("color"):
            run.font.color.rgb = RGBColor.from_string(item["color"].lstrip("#"))
        if item.get("bg_color"):
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), item["bg_color"].lstrip("#"))
            run._element.get_or_add_rPr().append(shd)
        if item.get("font"):
            _set_run_font(run, item["font"])
        if item.get("size"):
            run.font.size = Pt(item["size"])


_ALIGN_VALS = ("left", "center", "right", "justify")
_ALIGN_TO_DOCX = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _parse_block_attrs(e: dict) -> dict:
    """解析可选块级属性 align/indent/heading_level：只收集 payload 中提供的键
    （后端仅应用已提供的键，未提供的段落级属性原样保留）；非法值抛 ValueError。
    heading_level: null=正文 / 1-3=Heading 2-4（与前端 tag-1 约定一致）。"""
    attrs = {}
    if "align" in e:
        v = e.get("align")
        if v is not None:
            if v not in _ALIGN_VALS:
                raise ValueError(f"align 非法：{v}")
            attrs["align"] = v
    if "indent" in e:
        try:
            v = int(e.get("indent"))
        except (TypeError, ValueError):
            raise ValueError("indent 必须是整数")
        if not 0 <= v <= 8:
            raise ValueError("indent 超出范围（0-8）")
        attrs["indent"] = v
    if "heading_level" in e:
        v = e.get("heading_level")
        if v is not None:
            try:
                v = int(v)
            except (TypeError, ValueError):
                raise ValueError("heading_level 必须是整数或 null")
            if not 1 <= v <= 3:
                raise ValueError("heading_level 超出范围（1-3）")
        attrs["heading_level"] = v
    return attrs


def _apply_block_attrs(doc, p: DocxParagraph, attrs: dict):
    """应用块级属性（attrs 经 _parse_block_attrs 校验，只含提供的键）。
    标题样式按 builtin 名查找（Heading 2-4 / Normal），用户文档缺样式时
    best-effort 跳过；缩进每级 600 twips（≈ 编辑器 40px/级）。"""
    if "heading_level" in attrs:
        hl = attrs["heading_level"]
        try:
            p.style = doc.styles["Normal" if hl is None else f"Heading {hl + 1}"]
        except Exception:
            pass  # 样式缺失/文档定制样式表时跳过，文本改动不受影响
    if "align" in attrs:
        p.alignment = _ALIGN_TO_DOCX[attrs["align"]]
    if "indent" in attrs:
        p.paragraph_format.left_indent = Pt(attrs["indent"] * 30)


@manager.route("/flow/<flow_id>/document/edit", methods=["POST"])  # noqa: F821
@login_required
async def edit_document(flow_id: str):
    """Word 式正文编辑：按 /files/<id>/content 返回的 para_index 定位 docx 段落，
    支持三类操作——edits 改写段落文本 / deletes 删除段落 / inserts 新增段落
    （after_para_index=-1 表示插到文档开头，否则插到该段之后），全部基于原
    para_index，定位成功后统一应用，存为新版本（source=manual_edit）。
    表格/图片为原子块，不可改写、不可删除。"""
    try:
        flow = _require_owner(_flow_dict(flow_id))
        if flow["status"] in FlowWorkflow.TERMINAL:
            return _err("流程已结束")
        body = await request.get_json(silent=True) or {}
        version_id = body.get("version_id") or flow["current_version_id"]
        if not version_id:
            return _err("流程暂无文件版本，无法编辑", 101)

        edits = body.get("edits") or []
        deletes = body.get("deletes") or []
        inserts = body.get("inserts") or []
        if not isinstance(edits, list) or not isinstance(deletes, list) or not isinstance(inserts, list):
            return _err("edits/deletes/inserts 必须是数组", 101)
        if not edits and not deletes and not inserts:
            return _err("没有需要保存的改动", 101)
        if len(edits) + len(deletes) + len(inserts) > 200:
            return _err("单次最多修改 200 处", 101)

        parsed_edits = []
        edit_indexes = set()
        for e in edits:
            if not isinstance(e, dict):
                return _err("edits 项格式非法", 101)
            try:
                para_index = int(e.get("para_index"))
            except (TypeError, ValueError):
                return _err("para_index 必须是整数", 101)
            new_text = _CTRL_CHARS.sub("", str(e.get("new_text") or "")).strip()
            if not new_text:
                return _err("段落内容不能为空", 101)
            if len(new_text) > 20000:
                return _err("单段内容不能超过 20000 字", 101)
            try:
                runs = _parse_runs(e.get("runs"))
                block_attrs = _parse_block_attrs(e)
            except ValueError as ve:
                return _err(f"段落 {para_index} 格式非法：{ve}", 101)
            # 前端 newText 为 trim 后文本而 runs 来自未 trim 的块文本，两侧 strip 后再比
            if runs and "".join(x["text"] for x in runs).strip() != new_text:
                return _err(f"段落 {para_index} runs 文本与 new_text 不一致", 101)
            parsed_edits.append((para_index, new_text, runs, block_attrs))
            edit_indexes.add(para_index)

        parsed_deletes = []
        seen_deletes = set()
        for d in deletes:
            try:
                idx = int(d)
            except (TypeError, ValueError):
                return _err("deletes 项必须是整数段落号", 101)
            if idx in seen_deletes:
                return _err(f"段落 {idx} 重复删除", 101)
            seen_deletes.add(idx)
            parsed_deletes.append(idx)
        overlap = edit_indexes & set(parsed_deletes)
        if overlap:
            return _err(f"段落 {sorted(overlap)} 不能同时修改和删除", 101)

        parsed_inserts = []
        for ins in inserts:
            if not isinstance(ins, dict):
                return _err("inserts 项格式非法", 101)
            try:
                after = int(ins.get("after_para_index"))
            except (TypeError, ValueError):
                return _err("after_para_index 必须是整数（-1 表示文档开头）", 101)
            new_text = _CTRL_CHARS.sub("", str(ins.get("new_text") or "")).strip()
            if not new_text:
                return _err("新段落内容不能为空", 101)
            if len(new_text) > 20000:
                return _err("新段落内容不能超过 20000 字", 101)
            try:
                runs = _parse_runs(ins.get("runs"))
                block_attrs = _parse_block_attrs(ins)
            except ValueError as ve:
                return _err(f"新段落格式非法：{ve}", 101)
            if runs and "".join(x["text"] for x in runs).strip() != new_text:
                return _err(f"新段落 runs 文本与 new_text 不一致", 101)
            parsed_inserts.append((after, new_text, runs, block_attrs))

        version = next(
            (v for v in FlowVersionService.list_by_flow(flow_id) if v["id"] == version_id), None
        )
        if not version:
            return _err("版本不存在", 404)
        # 只允许编辑当前版本：编辑历史版本会以 add_version 落成最新版，
        # 等于把文件内容悄悄回滚，绕过「文档已变化请刷新」保护
        if version_id != flow["current_version_id"]:
            return _err("文档已有新版本，请刷新后重试", 101)

        blob = await thread_pool_exec(settings.STORAGE_IMPL.get, _bucket_of(flow), version["file_path"])
        if not blob:
            return _err("版本文件读取失败", 101)

        # docx 直接编辑；老格式 .doc 先经 LibreOffice 转 docx（与 /files/<id>/content
        # 的解析路径一致，para_index 映射保持同源），编辑后新版本统一存为 .docx
        is_docx = (version["file_name"] or "").lower().endswith(".docx")
        if not is_docx:
            if not is_doc_file(blob, version["file_name"] or ""):
                return _err("仅支持编辑 doc/docx 文档", 101)
            blob = await thread_pool_exec(doc_to_docx_via_libreoffice, blob)
            if not blob:
                return _err(".doc 转换失败，暂无法编辑该文档", 101)

        doc = await thread_pool_exec(DocxDocument, BytesIO(blob))
        para_map = await thread_pool_exec(_build_para_map, doc)
        keys_sorted = sorted(para_map)

        def _entry_of(idx):
            entry = para_map.get(idx)
            if not entry:
                return None
            kind, target = entry
            if kind != "p":
                label = "表格" if kind == "table" else "图片"
                return ("atomic", label, None)
            return ("p", None, target)

        # 先全部定位成功，再统一应用（避免半改状态）
        located_edits = []
        for para_index, new_text, runs, block_attrs in parsed_edits:
            entry = _entry_of(para_index)
            if entry is None:
                return _err(f"段落 {para_index} 定位失败，文档可能已变化，请刷新后重试", 101)
            if entry[0] == "atomic":
                return _err(f"段落 {para_index} 是{entry[1]}，不支持编辑", 101)
            located_edits.append((entry[2], new_text, runs, block_attrs))

        located_deletes = []
        for para_index in parsed_deletes:
            entry = _entry_of(para_index)
            if entry is None:
                return _err(f"段落 {para_index} 定位失败，文档可能已变化，请刷新后重试", 101)
            if entry[0] == "atomic":
                return _err(f"段落 {para_index} 是{entry[1]}，不支持删除", 101)
            located_deletes.append(entry[2])

        located_inserts = []  # (mode, ref_paragraph|None, text, style_src|None, runs|None, block_attrs|None)
        delete_set = set(parsed_deletes)
        for after, new_text, runs, block_attrs in parsed_inserts:
            if after >= 0 and para_map.get(after) is None and after != -1:
                return _err(f"插入位置 {after} 无效", 101)
            style_src = None
            if after >= 0:
                anchor = para_map.get(after)
                if anchor and anchor[0] == "p":
                    style_src = anchor[1]
            # 插入参照段必须跳过本请求将删除的段落：ref 先被删会脱离文档树，
            # insert_paragraph_before 对已分离元素静默写入空气段（内容丢失）
            nxt = next(
                (
                    para_map[k][1]
                    for k in keys_sorted
                    if k > after and k not in delete_set and para_map[k][0] == "p"
                ),
                None,
            )
            if nxt is not None:
                located_inserts.append(("before", nxt, new_text, style_src, runs, block_attrs))
            else:
                located_inserts.append(("append", None, new_text, style_src, runs, block_attrs))

        # 应用顺序：删除 → 插入 → 改写（改写持有元素引用，不受结构变化影响）；
        # 逐 run 的 oxml 操作是 CPU 密集，放线程池避免阻塞事件循环
        def _apply_ops():
            for p in located_deletes:
                el = p._element
                el.getparent().remove(el)
            for mode, ref, text, style_src, rns, attrs in located_inserts:
                if mode == "append":
                    new_p = doc.add_paragraph(text)
                else:
                    new_p = ref.insert_paragraph_before(text)
                if style_src is not None:
                    try:
                        new_p.style = style_src.style
                    except Exception:
                        pass
                if attrs:
                    _apply_block_attrs(doc, new_p, attrs)
                if rns is not None:
                    _apply_runs(new_p, rns)
            for target, text, rns, attrs in located_edits:
                if attrs:
                    _apply_block_attrs(doc, target, attrs)
                if rns is not None:
                    _apply_runs(target, rns)
                else:
                    _replace_para_text(target, text)

        await thread_pool_exec(_apply_ops)

        out = BytesIO()
        await thread_pool_exec(doc.save, out)
        new_blob = out.getvalue()

        root, _ = os.path.splitext(version["file_name"] or "document")
        file_name = f"{_safe_filename(root)}.docx" if not is_docx else f"{_safe_filename(root)}_edited.docx"
        no = FlowVersionService.next_version_no(flow_id)
        object_name = f"flow/{flow_id}/v{no}_{file_name}"
        await thread_pool_exec(settings.STORAGE_IMPL.put, _bucket_of(flow), object_name, new_blob)
        new_version = FlowVersionService.add_version(
            flow, object_name, file_name,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            len(new_blob), "manual_edit", current_user.id,
        )
        return get_json_result(data={"version": new_version})
    except LookupError as e:
        return _err(str(e), 404)
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 5. 下载版本文件 ───────────────────────────────────────────────
@manager.route("/flow/<flow_id>/version/<version_id>/download", methods=["GET"])  # noqa: F821
@login_required
async def download_version(flow_id: str, version_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        version = next((v for v in FlowVersionService.list_by_flow(flow_id) if v["id"] == version_id), None)
        if not version:
            return _err("版本不存在", 404)
        blob = await thread_pool_exec(settings.STORAGE_IMPL.get, _bucket_of(flow), version["file_path"])
        file_name = version["file_name"] or "file"
        return Response(
            blob,
            mimetype=version["file_type"] or "application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(file_name)}"},
        )
    except LookupError as e:
        return _err(str(e), 404)
    except PermissionError as e:
        return _err(str(e), 403)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 6. 添加批注 ───────────────────────────────────────────────────
@manager.route("/flow/<flow_id>/comment", methods=["POST"])  # noqa: F821
@login_required
async def add_comment(flow_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        if flow["status"] in FlowWorkflow.TERMINAL:
            return _err("流程已结束")
        body = await request.get_json(silent=True) or {}
        content = (body.get("content") or "").strip()
        if not content:
            return _err("批注内容不能为空", 101)
        version_id = body.get("version_id") or flow["current_version_id"]
        if not version_id:
            return _err("流程暂无文件版本，无法批注", 101)
        # Word 式批注锚点：选中的原文选段 + 段落 index + 段落内起始偏移（均可空）
        anchor_text = (body.get("anchor_text") or "").strip()[:500]
        anchor_para = body.get("anchor_para")
        if anchor_para is not None and not isinstance(anchor_para, int):
            try:
                anchor_para = int(anchor_para)
            except (TypeError, ValueError):
                anchor_para = None
        anchor_start = body.get("anchor_start")
        if anchor_start is not None and not isinstance(anchor_start, int):
            try:
                anchor_start = int(anchor_start)
            except (TypeError, ValueError):
                anchor_start = None

        comment = FlowCommentService.add_comment(
            flow_id, version_id, current_user.id, content,
            anchor_text=anchor_text, anchor_para=anchor_para,
            anchor_start=anchor_start,
        )
        others = _others_of(flow, current_user.id)
        try:
            notify_flow_event(
                flow, others,
                f"流程「{flow['title']}」有新批注",
                f"{_nickname_of(current_user.id)} 添加了批注意见",
            )
        except Exception as e:
            logger.warning("flow notify failed: %s", e)
        return get_json_result(data={"comment": comment})
    except LookupError as e:
        return _err(str(e), 404)
    except PermissionError as e:
        return _err(str(e), 403)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 6.1 删除批注（仅批注作者本人） ────────────────────────────────
@manager.route("/flow/<flow_id>/comment/<comment_id>/delete", methods=["POST"])  # noqa: F821
@login_required
async def delete_comment(flow_id: str, comment_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        if flow["status"] in FlowWorkflow.TERMINAL:
            return _err("流程已结束")
        comment = FlowCommentService.get_comment(comment_id)
        if not comment or comment["flow_id"] != flow_id:
            return _err("批注不存在", 404)
        if comment["user_id"] != current_user.id:
            return _err("只能删除自己的批注", 403)
        FlowCommentService.delete_comment(comment_id)
        return get_json_result(data={"id": comment_id})
    except LookupError as e:
        return _err(str(e), 404)
    except PermissionError as e:
        return _err(str(e), 403)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 7. AI 处理记录 ────────────────────────────────────────────────
@manager.route("/flow/<flow_id>/ai-record", methods=["POST"])  # noqa: F821
@login_required
async def add_ai_record(flow_id: str):
    try:
        flow = _require_owner(_flow_dict(flow_id))
        body = await request.get_json(silent=True) or {}
        instruction = (body.get("instruction") or "").strip()
        response = (body.get("response") or "").strip()
        if not instruction:
            return _err("指令内容不能为空", 101)
        if not response:
            return _err("AI 回复内容不能为空", 101)
        version_id = body.get("version_id") or flow["current_version_id"]
        if not version_id:
            return _err("流程暂无文件版本，无法记录 AI 处理", 101)
        session_id = body.get("session_id") or ""
        save_as_version = body.get("save_as_version", True)

        output_version_id = ""
        if save_as_version:
            # 秒级时间戳同秒多次保存会撞名覆盖，追加 8 位随机后缀兜底
            object_name = f"flow/{flow_id}/ai_{int(time.time())}_{uuid.uuid4().hex[:8]}.md"
            blob = response.encode("utf-8")
            await thread_pool_exec(settings.STORAGE_IMPL.put, _bucket_of(flow), object_name, blob)
            version = FlowVersionService.add_version(
                flow, object_name,
                f"AI产出_{time.strftime('%Y%m%d_%H%M%S')}.md",
                "text/markdown", len(blob), "ai_output", current_user.id,
            )
            output_version_id = version["id"]

        record = FlowAiChatService.add_record(
            flow_id, version_id, instruction, response, session_id, output_version_id,
        )
        return get_json_result(data={"record": record, "output_version_id": output_version_id})
    except LookupError as e:
        return _err(str(e), 404)
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 8. 流转 / 退回 ────────────────────────────────────────────────
@manager.route("/flow/<flow_id>/submit", methods=["POST"])  # noqa: F821
@login_required
async def submit_flow(flow_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        body = await request.get_json(silent=True) or {}
        action = body.get("action")
        if action not in ("next", "return"):
            return _err("action 必须是 next 或 return", 101)
        updated = FlowActionService.submit(flow, current_user.id, action)
        try:
            notify_target_of(updated, action)
        except Exception as e:
            logger.warning("flow notify failed: %s", e)
        return get_json_result(data={"flow": updated})
    except LookupError as e:
        return _err(str(e), 404)
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 9. 归档 ───────────────────────────────────────────────────────
@manager.route("/flow/<flow_id>/archive", methods=["POST"])  # noqa: F821
@login_required
async def archive_flow(flow_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        updated = FlowActionService.archive(flow, current_user.id)
        try:
            notify_target_of(updated, "archive")
        except Exception as e:
            logger.warning("flow notify failed: %s", e)
        return get_json_result(data={"flow": updated})
    except LookupError as e:
        return _err(str(e), 404)
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))


# ── 10. 作废 ──────────────────────────────────────────────────────
@manager.route("/flow/<flow_id>/cancel", methods=["POST"])  # noqa: F821
@login_required
async def cancel_flow(flow_id: str):
    try:
        flow = _require_participant(_flow_dict(flow_id))
        updated = FlowActionService.cancel(flow, current_user.id)
        try:
            notify_target_of(updated, "cancel")
        except Exception as e:
            logger.warning("flow notify failed: %s", e)
        return get_json_result(data={"flow": updated})
    except LookupError as e:
        return _err(str(e), 404)
    except (PermissionError, ValueError, RuntimeError) as e:
        return _action_error(e)
    except Exception as e:
        logger.exception(e)
        return _err(str(e))
