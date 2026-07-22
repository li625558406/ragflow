#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import base64
import html
import io
import logging
import os
import re

from werkzeug.security import generate_password_hash, check_password_hash

from common import settings
from api.db import TenantPermission
from api.db.db_models import (
    CollaborationAttachment,
    CollaborationAuditLog,
    CollaborationComment,
    CollaborationDocument,
    CollaborationDocumentVersion,
    User,
    CollaborationDocumentACL,
    CollaborationFolder,
    CollaborationFormatRule,
    CollaborationShareLink,
    DB,
    UserTenant,
)
from api.db.services.collaboration_service import (
    CollaborationAttachmentService,
    CollaborationAuditLogService,
    CollaborationCommentService,
    CollaborationDocumentACLService,
    CollaborationDocumentService,
    CollaborationDocumentVersionService,
    CollaborationFolderService,
    CollaborationFormatRuleService,
    CollaborationShareLinkService,
)
from api.db.services.user_service import UserTenantService
from common.misc_utils import get_uuid
from common.time_utils import current_timestamp

# 历史版本保留条数：超过则删除最旧的。
MAX_DOCUMENT_VERSIONS_KEEP = 20


def _get_shared_tenant_user_ids(user_id: str) -> set:
    """Get the set of user IDs that share at least one tenant with the given user.

    In RAGFlow, team membership is modeled via UserTenant (user_id → tenant_id).
    Two users are on the same team if they share at least one tenant_id.
    """
    # Tenant IDs that the given user belongs to
    user_records = UserTenantService.query(user_id=user_id)
    tenant_ids = [t.tenant_id for t in user_records]
    if not tenant_ids:
        return set()
    # All user IDs belonging to those tenants (including the querying user)
    members = UserTenant.select(UserTenant.user_id).where(
        UserTenant.tenant_id.in_(tenant_ids)
    )
    return {m.user_id for m in members}


# ── Security Utilities ──

def sanitize_html(text: str) -> str:
    """Escape HTML + strip dangerous attributes. For comment/document content."""
    text = html.escape(text, quote=True)
    text = re.sub(r'javascript\s*:', '', text, flags=re.IGNORECASE)
    return text


def sanitize_filename(filename: str) -> str:
    """Keep only safe chars, prevent path traversal."""
    name = os.path.basename(filename)
    name = re.sub(r'[^\w\.\-]', '_', name)
    return name or 'untitled'


ALLOWED_UPLOAD_MIMES = {
    'image/jpeg', 'image/png', 'image/gif', 'image/webp',
    'application/pdf', 'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'text/plain', 'text/csv',
}
MAX_UPLOAD_SIZE = 50 * 1024 * 1024  # 50MB


def validate_upload(file_obj) -> tuple[bytes, str, str]:
    """Validate uploaded file: MIME whitelist, size limit, filename sanitization.

    Returns (data, safe_filename, mime_type).
    Raises ValueError on validation failure.
    """
    mime_type = file_obj.content_type or 'application/octet-stream'
    if mime_type not in ALLOWED_UPLOAD_MIMES:
        raise ValueError(f"Unsupported file type: {mime_type}")

    # Pre-check Content-Length to reject oversize files before reading
    cl = file_obj.content_length
    if cl is not None and cl > MAX_UPLOAD_SIZE:
        raise ValueError(f"File too large (max 50MB)")

    data = file_obj.read()
    if len(data) > MAX_UPLOAD_SIZE:
        raise ValueError(f"File too large (max 50MB)")

    safe_name = sanitize_filename(file_obj.filename or 'untitled')
    return data, safe_name, mime_type


async def log_audit(user_id: str, document_id: str | None, action: str, detail: dict = None, ip_address: str = None):
    """Record an audit log entry. Fire-and-forget — failures are logged but not raised."""
    try:
        CollaborationAuditLogService.save(
            id=get_uuid(),
            user_id=user_id,
            document_id=document_id,
            action=action,
            detail=detail or {},
            ip_address=ip_address,
            create_time=current_timestamp(),
        )
    except Exception as e:
        logging.warning(f"Audit log failed for {action}: {e}")
ROLE_HIERARCHY = {
    "owner": 4,
    "editor": 3,
    "viewer": 2,
    "commenter": 1,
}


def _get_user_role(doc_id: str, user_id: str) -> str | None:
    """Get the effective role of a user on a document, or None if no access.

    Priority: ACL entry > owner by created_by > team permission fallback
    """
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        return None

    # Owner always has owner role
    if doc.created_by == user_id:
        return "owner"

    # Check ACL for explicit grant
    acl = CollaborationDocumentACLService.query(document_id=doc_id, user_id=user_id)
    if acl:
        return acl[0].role

    # Team permission fallback: team-visible docs → editor role (full CRUD)
    if doc.permission == TenantPermission.TEAM.value:
        team_user_ids = _get_shared_tenant_user_ids(user_id)
        if doc.created_by in team_user_ids:
            return "editor"

    return None


def _check_access(obj, user_id: str) -> bool:
    """Check if user has access to a document/format-rule/folder, following agent team permission model.

    For objects with a 'permission' field (documents, format rules), respects the permission setting.
    For objects without a 'permission' field (folders), allows team members access.
    """
    if obj.created_by == user_id:
        return True
    perm = getattr(obj, 'permission', None)
    if perm is not None and perm != TenantPermission.TEAM.value:
        return False
    team_user_ids = _get_shared_tenant_user_ids(user_id)
    return obj.created_by in team_user_ids


def _check_role(doc_id: str, user_id: str, required_role: str) -> bool:
    """Check if user has at least the required role on a document."""
    role = _get_user_role(doc_id, user_id)
    if role is None:
        return False
    return ROLE_HIERARCHY.get(role, 0) >= ROLE_HIERARCHY.get(required_role, 0)


async def create_document(tenant_id: str, user_id: str, name: str, markdown_content: str, agent_id: str = None, permission: str = "me", folder_id: str = None) -> dict:
    """Create a collaboration document from chat message content.

    The content field is seeded with a minimal Univer Docs skeleton; the
    Univer frontend's createDocument loader fills in any missing fields
    (id/rev/style) on first open. The original markdown is preserved in
    markdown_content for the frontend to display the source message and
    for a future markdown->Univer migration path.
    """
    doc_id = get_uuid()
    # Univer Docs blank skeleton — createDocument on the frontend completes
    # the remaining IDocumentData fields.
    content = {"document": True, "body": {"blockType": "paragraph", "children": []}}
    CollaborationDocumentService.save(
        id=doc_id,
        name=name,
        file_type="docx",
        content=content,
        markdown_content=markdown_content,
        tenant_id=tenant_id,
        created_by=user_id,
        agent_id=agent_id,
        permission=permission,
        folder_id=folder_id,
    )
    return {"id": doc_id, "name": name, "file_type": "docx", "permission": permission, "folder_id": folder_id}


async def create_spreadsheet(tenant_id: str, user_id: str, name: str, permission: str = "me", folder_id: str = None) -> dict:
    """Create a blank collaboration spreadsheet document."""
    doc_id = get_uuid()
    default_content = {
        "sheets": [{"name": "Sheet1", "data": [[""]], "colWidths": [100]}],
        "activeSheet": 0,
    }
    CollaborationDocumentService.save(
        id=doc_id,
        name=name,
        file_type="xlsx",
        content=default_content,
        markdown_content="",
        tenant_id=tenant_id,
        created_by=user_id,
        permission=permission,
        folder_id=folder_id,
    )
    return {"id": doc_id, "name": name, "file_type": "xlsx", "permission": permission, "folder_id": folder_id}


def _generate_xlsx(content: dict) -> bytes:
    """Generate .xlsx bytes from spreadsheet content.

    Supports both the new native IWorkbookData format (preferred) and the
    legacy `{sheets:[{name,data,colWidths}]}` grid format for documents
    created before the M1 spreadsheet upgrade. Conversion is delegated to
    spreadsheet_xlsx_adapter so all openpyxl↔Univer mapping lives in one place.
    """
    from api.apps.services.spreadsheet_xlsx_adapter import workbook_data_to_xlsx

    return workbook_data_to_xlsx(content or {})


async def import_xlsx(tenant_id: str, user_id: str, file_obj, folder_id: str = None) -> dict:
    """Parse a .xlsx file and create a collaboration spreadsheet document.

    Storage format: native Univer IWorkbookData (since M1). Legacy documents
    continue to render via the frontend's convertLegacyToWorkbookData fallback.
    """
    from api.apps.services.spreadsheet_xlsx_adapter import (
        XlsxTooLargeError,
        xlsx_to_workbook_data,
    )

    doc_id = get_uuid()
    name = (file_obj.filename or "imported").rsplit(".", 1)[0]

    data = file_obj.read()
    # xlsx_to_workbook_data enforces MAX_XLSX_SIZE / MAX_ROWS / MAX_COLS
    # and raises XlsxTooLargeError (a ValueError subclass) — the API layer
    # maps that to a 4xx for the frontend to render a friendly Toast.
    content = xlsx_to_workbook_data(data, doc_id)

    CollaborationDocumentService.save(
        id=doc_id,
        name=name,
        file_type="xlsx",
        content=content,
        markdown_content="",
        tenant_id=tenant_id,
        created_by=user_id,
        permission="me",
        folder_id=folder_id,
    )
    await log_audit(tenant_id, doc_id, "document.import_xlsx", {"name": name})

    return {"id": doc_id, "name": name, "file_type": "xlsx", "folder_id": folder_id}


async def import_docx(tenant_id: str, user_id: str, file_obj, folder_id: str = None) -> dict:
    """Parse a .docx file → Univer Docs IDocumentData snapshot.

    python-docx caveat: `ParagraphFormat.outline_level` does NOT exist in the
    public API (it raises AttributeError). Heading level must be detected via
    `para.style.name` ("Heading 1" / "Heading 2" / ...) — that is the only
    stable public path.

    Output: Univer IDocumentData with body.dataStream / textRuns / paragraphs.
    Tables are rendered as tab-separated rows (Univer Docs preset we loaded
    doesn't include a native table block).
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx not installed")

    try:
        doc = DocxDocument(file_obj)
    except Exception as e:
        err_msg = str(e).lower()
        if "not a word file" in err_msg or "spreadsheet" in err_msg:
            raise ValueError(
                "该文件不是 Word 文档(.docx)。如果要导入 Excel 文件，请使用 .xlsx 扩展名。"
            )
        raise

    doc_id = get_uuid()
    name = (file_obj.filename or "imported").rsplit(".", 1)[0]

    # Univer IDocumentBody dataStream character semantics (see
    # @univerjs/core i-document-data.d.ts):
    #   \r  PARAGRAPH    (paragraph terminator)
    #   \n  SECTION_BREAK (section break — NOT a paragraph separator!)
    #   \t  TAB
    #   \v  COLUMN_BREAK
    #   \f  PAGE_BREAK
    #
    # Earlier version of this code used "\r\n" between every paragraph,
    # which inserted an unintended SECTION_BREAK after each paragraph and
    # shattered the document into N empty sections → Univer rendered blank.
    # Correct pattern: terminate each paragraph with "\r", and add a single
    # trailing "\n" at end-of-document so the body has one valid section.
    #
    # paragraphStyle uses namedStyleType (NamedStyleType enum):
    #   HEADING_1=4, HEADING_2=5, HEADING_3=6, HEADING_4=7, HEADING_5=8
    # There is NO `headingLevel` field on IParagraphStyle.
    HEADING_TO_NAMED = {1: 4, 2: 5, 3: 6, 4: 7, 5: 8, 6: 9}

    data_stream_parts: list[str] = []
    text_runs: list[dict] = []
    paragraphs_meta: list[dict] = []
    markdown_lines: list[str] = []
    cursor = 0

    def _emit_runs(runs) -> int:
        nonlocal cursor
        added = 0
        for run in runs:
            run_text = run.text or ""
            # Strip Univer control chars that would create unintended
            # paragraph / section / page / column breaks inside a run.
            # Common source: .docx soft line breaks (<w:br/>) surface as
            # "\n" inside run.text — Univer interprets that as SECTION_BREAK
            # and splits the doc mid-paragraph, producing phantom blank
            # pages. Tabs are also stripped (Univer treats "\t" specially).
            for ch in ("\r", "\n", "\v", "\f", "\t"):
                run_text = run_text.replace(ch, "")
            if not run_text:
                continue
            ts: dict = {}
            if run.bold:
                ts["bl"] = 1
            if run.italic:
                ts["it"] = 1
            if run.underline:
                ts["ul"] = {"s": 1}
            if run.font and run.font.strike:
                ts["st"] = {"s": 1}
            text_runs.append({
                "st": cursor,
                "ed": cursor + len(run_text),
                "ts": ts,
            })
            data_stream_parts.append(run_text)
            cursor += len(run_text)
            added += len(run_text)
        return added

    def _close_paragraph(para_start: int, named_style_type: int | None) -> None:
        nonlocal cursor
        paragraph_style: dict = {}
        if named_style_type is not None:
            paragraph_style["namedStyleType"] = named_style_type
        paragraphs_meta.append({
            "startIndex": para_start,
            "paragraphStyle": paragraph_style,
        })
        data_stream_parts.append("\r")
        cursor += 1

    # Track whether we've emitted any real content yet, to skip leading
    # empty paragraphs (which would push the first heading to page 2).
    # Also collapse consecutive empty paragraphs into one to avoid stacks
    # of blank lines (a common artifact of .docx export tools).
    seen_content = False
    last_was_empty = False

    for para in doc.paragraphs:
        # Heading detection via style name (the only public API path).
        # DO NOT use para.paragraph_format.outline_level — that attribute
        # does not exist and will throw AttributeError mid-import.
        heading_level: int | None = None
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if style_name == "title" or "heading 1" in style_name:
                heading_level = 1
            elif "heading 2" in style_name:
                heading_level = 2
            elif "heading 3" in style_name:
                heading_level = 3
            elif "heading 4" in style_name:
                heading_level = 4
            elif "heading 5" in style_name:
                heading_level = 5
            elif "heading 6" in style_name:
                heading_level = 6

        para_start = cursor
        added = _emit_runs(para.runs)
        # Fallback: if runs produced nothing but para.text has content (e.g.
        # field codes, form fields), emit the visible text as a single run.
        # Same control-char sanitization as _emit_runs.
        if added == 0 and para.text:
            fallback_text = para.text
            for ch in ("\r", "\n", "\v", "\f", "\t"):
                fallback_text = fallback_text.replace(ch, "")
            if fallback_text:
                text_runs.append({
                    "st": cursor,
                    "ed": cursor + len(fallback_text),
                    "ts": {},
                })
                data_stream_parts.append(fallback_text)
                cursor += len(fallback_text)
                added = len(fallback_text)

        is_empty = added == 0
        # Skip leading empty paragraphs (before any content) and collapse
        # consecutive empties into one. Trailing empties are stripped by
        # the "last_was_empty && is_empty" rule and the absence of further
        # content.
        if is_empty:
            if not seen_content:
                continue  # leading blank — drop
            if last_was_empty:
                continue  # collapse stack of blanks — keep only one
            last_was_empty = True
        else:
            seen_content = True
            last_was_empty = False

        # Markdown mirror (used for full-text search / list preview)
        text_md = para.text.strip()
        if heading_level and text_md:
            markdown_lines.append("#" * heading_level + " " + text_md)
        else:
            markdown_lines.append(text_md)

        _close_paragraph(para_start, HEADING_TO_NAMED.get(heading_level))

    # Tables → one paragraph per row, cells joined by tab.
    for table in doc.tables:
        for row in table.rows:
            row_start = cursor
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            row_text = "\t".join(cells)
            if row_text:
                text_runs.append({
                    "st": row_start,
                    "ed": row_start + len(row_text),
                    "ts": {},
                })
                data_stream_parts.append(row_text)
                cursor += len(row_text)
            _close_paragraph(row_start, None)
            markdown_lines.append("| " + " | ".join(cells) + " |")

    # Body must contain at least one paragraph or Univer's renderer throws.
    # Minimal valid body = empty paragraph + section break.
    if not data_stream_parts:
        paragraphs_meta.append({"startIndex": 0, "paragraphStyle": {}})
        data_stream_parts.append("\r")

    # Trailing section break — Univer expects body to end with a section.
    data_stream_parts.append("\n")

    content = {
        "id": "default_doc",
        "documentStyle": {
            "pageSize": {"width": 794, "height": 1124},  # A4 @ 96dpi
            "documentFlavor": 1,  # DocumentFlavor.TRADITIONAL (A4 paged)
            "marginTop": 50,
            "marginBottom": 50,
            "marginLeft": 90,
            "marginRight": 90,
        },
        "body": {
            "dataStream": "".join(data_stream_parts),
            "textRuns": text_runs,
            "paragraphs": paragraphs_meta,
        },
    }
    markdown_content = "\n".join(markdown_lines)

    CollaborationDocumentService.save(
        id=doc_id,
        name=name,
        file_type="docx",
        folder_id=folder_id,
        content=content,
        markdown_content=markdown_content,
        tenant_id=tenant_id,
        created_by=user_id,
        permission="me",
    )
    await log_audit(tenant_id, doc_id, "document.import_docx", {"name": name})

    return {"id": doc_id, "name": name, "file_type": "docx", "folder_id": folder_id}


async def list_documents(tenant_id: str, user_id: str) -> list:
    """List collaboration documents visible to the current user (own + team-shared + ACL)."""
    team_user_ids = _get_shared_tenant_user_ids(user_id)
    # Get document IDs where user has explicit ACL grant
    acl_doc_ids = [
        a.document_id
        for a in CollaborationDocumentACLService.query(user_id=user_id)
    ]
    base_condition = (
        (
            (CollaborationDocument.created_by.in_(team_user_ids))
            & (CollaborationDocument.permission == TenantPermission.TEAM.value)
        )
        | (CollaborationDocument.created_by == user_id)
    )
    if acl_doc_ids:
        base_condition |= CollaborationDocument.id.in_(acl_doc_ids)
    docs = (
        CollaborationDocument.select()
        .where(base_condition)
        .order_by(CollaborationDocument.create_time.desc())
    )
    result = []
    for d in docs:
        result.append({
            "id": d.id,
            "name": d.name,
            "file_type": d.file_type,
            "agent_id": d.agent_id,
            "folder_id": d.folder_id,
            "sort_order": d.sort_order,
            "create_time": d.create_time,
            "update_time": d.update_time,
            "created_by": d.created_by,
            "permission": d.permission,
        })
    return result


async def get_document(doc_id: str, tenant_id: str) -> dict:
    """Get a single collaboration document with content."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    user_role = _get_user_role(doc_id, tenant_id)
    if not user_role:
        raise PermissionError("Access denied")
    return {
        "id": doc.id,
        "name": doc.name,
        "file_type": doc.file_type,
        "file_path": doc.file_path,
        "content": doc.content,
        "markdown_content": doc.markdown_content,
        "agent_id": doc.agent_id,
        "created_by": doc.created_by,
        "permission": doc.permission,
        "role": user_role,
        "create_time": doc.create_time,
        "update_time": doc.update_time,
        "ydoc": base64.b64encode(doc.ydoc).decode("ascii") if doc.ydoc else None,
        "version": doc.version or 0,
    }


async def update_document(doc_id: str, tenant_id: str, data: dict) -> dict:
    """Update document name and/or content. Requires editor role or higher."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "editor"):
        raise PermissionError("Access denied")

    update_data = {}
    if "name" in data:
        update_data["name"] = data["name"]
    if "content" in data:
        update_data["content"] = data["content"]
    if "markdown_content" in data:
        update_data["markdown_content"] = data["markdown_content"]
        update_data["file_path"] = None  # invalidate cached file so download regenerates
    if "permission" in data:
        update_data["permission"] = data["permission"]
    if "folder_id" in data:
        update_data["folder_id"] = data["folder_id"]
    if "sort_order" in data:
        update_data["sort_order"] = data["sort_order"]
    if "ydoc_state" in data:
        update_data["ydoc"] = base64.b64decode(data["ydoc_state"])
        update_data["version"] = (doc.version or 0) + 1

    if update_data:
        CollaborationDocumentService.update_by_id(doc_id, update_data)
        await log_audit(tenant_id, doc_id, "document.update", {"fields": list(update_data.keys())})
    return {"id": doc_id, "updated": list(update_data.keys())}


async def save_ydoc_state(doc_id: str, tenant_id: str, data: dict) -> dict:
    """Save Yjs binary state from the frontend (periodic persistence).

    每次内容保存都会在 CollaborationDocumentVersion 表里写一条快照，
    保留最新 MAX_DOCUMENT_VERSIONS_KEEP 条，用于真正的历史版本回滚。
    """
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "editor"):
        raise PermissionError("Access denied")

    update_data = {}
    ydoc_state_b64 = data.get("ydoc_state")
    new_version = None
    if ydoc_state_b64:
        update_data["ydoc"] = base64.b64decode(ydoc_state_b64)
        new_version = (doc.version or 0) + 1
        update_data["version"] = new_version
    if "content" in data:
        update_data["content"] = data["content"]
    if "markdown_content" in data:
        update_data["markdown_content"] = data["markdown_content"]
        update_data["file_path"] = None

    if update_data:
        CollaborationDocumentService.update_by_id(doc_id, update_data)

    # 写入版本快照（只在确实有 ydoc_state 时才记，元数据修改不入版本库）。
    if new_version is not None:
        try:
            snapshot_id = get_uuid()
            now_ms = current_timestamp()
            CollaborationDocumentVersionService.save(
                id=snapshot_id,
                document_id=doc_id,
                version=new_version,
                ydoc_snapshot=update_data.get("ydoc"),
                content_snapshot=update_data.get("content"),
                created_by=tenant_id,
                create_time=now_ms,
            )
            _trim_document_versions(doc_id)
        except Exception as exc:
            # 快照失败不能影响保存主流程。
            logging.warning(f"[save_ydoc_state] snapshot failed for doc {doc_id}: {exc}")

    return {"id": doc_id, "version": new_version if new_version is not None else (doc.version or 0)}


def _trim_document_versions(doc_id: str) -> None:
    """保留最新 MAX_DOCUMENT_VERSIONS_KEEP 条快照，其余删除。"""
    try:
        keep_ids = (
            CollaborationDocumentVersion
            .select(CollaborationDocumentVersion.id)
            .where(CollaborationDocumentVersion.document_id == doc_id)
            .order_by(
                CollaborationDocumentVersion.create_time.desc(),
                CollaborationDocumentVersion.version.desc(),
            )
            .limit(MAX_DOCUMENT_VERSIONS_KEEP)
        )
        CollaborationDocumentVersion.delete().where(
            (CollaborationDocumentVersion.document_id == doc_id)
            & (~(CollaborationDocumentVersion.id.in_(keep_ids)))
        ).execute()
    except Exception as exc:
        logging.warning(f"[save_ydoc_state] trim failed for doc {doc_id}: {exc}")


async def delete_document(doc_id: str, tenant_id: str) -> bool:
    """Delete a collaboration document and its stored file. Requires owner role."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")

    # Delete stored file if exists
    if doc.file_path:
        try:
            settings.STORAGE_IMPL.rm("collaboration", doc.file_path)
        except Exception as ex:
            logging.warning(f"Failed to delete file {doc.file_path}: {ex}")

    CollaborationDocumentService.delete_by_id(doc_id)
    await log_audit(tenant_id, doc_id, "document.delete", {"name": doc.name})
    return True


async def download_document(doc_id: str, tenant_id: str, file_type: str = "docx") -> tuple:
    """下载最近一次导出的文件。不再在服务端生成 docx/pdf。

    前端 Univer Docs SDK 导出后通过 /exported-file 端点上传，
    后端只存文件不生成。返回 None 表示尚未导出过。
    """
    return await get_exported_file(doc_id, tenant_id)


# ── Folder CRUD ──

async def create_folder(tenant_id: str, user_id: str, name: str, parent_id: str = None) -> dict:
    """Create a collaboration folder."""
    folder_id = get_uuid()
    CollaborationFolderService.save(
        id=folder_id,
        name=name,
        parent_id=parent_id,
        tenant_id=tenant_id,
        created_by=user_id,
        create_time=current_timestamp(),
    )
    return {"id": folder_id, "name": name, "parent_id": parent_id}


async def list_folders(tenant_id: str, user_id: str) -> list:
    """List folders visible to the current user, returned as flat list with tree data."""
    team_user_ids = _get_shared_tenant_user_ids(user_id)
    folders = (
        CollaborationFolder.select()
        .where(CollaborationFolder.created_by.in_(team_user_ids))
        .order_by(CollaborationFolder.sort_order)
    )
    return [
        {
            "id": f.id,
            "name": f.name,
            "parent_id": f.parent_id,
            "created_by": f.created_by,
            "sort_order": f.sort_order,
            "create_time": f.create_time,
        }
        for f in folders
    ]


async def update_folder(folder_id: str, tenant_id: str, data: dict) -> dict:
    """Update folder name or parent."""
    e, folder = CollaborationFolderService.get_by_id(folder_id)
    if not e:
        raise LookupError("Folder not found")
    if not _check_access(folder, tenant_id):
        raise PermissionError("Access denied")
    update_data = {}
    for key in ("name", "parent_id", "sort_order"):
        if key in data:
            update_data[key] = data[key]
    if update_data:
        CollaborationFolderService.update_by_id(folder_id, update_data)
    return {"id": folder_id, "updated": list(update_data.keys())}


async def delete_folder(folder_id: str, tenant_id: str) -> bool:
    """Delete a folder. Documents in the folder become root-level."""
    e, folder = CollaborationFolderService.get_by_id(folder_id)
    if not e:
        raise LookupError("Folder not found")
    if not _check_access(folder, tenant_id):
        raise PermissionError("Access denied")
    # Move child documents to root and child folders to root atomically
    with DB.atomic():
        CollaborationDocument.update(folder_id=None).where(
            CollaborationDocument.folder_id == folder_id
        ).execute()
        CollaborationFolder.update(parent_id=None).where(
            CollaborationFolder.parent_id == folder_id
        ).execute()
        CollaborationFolderService.delete_by_id(folder_id)
    return True


async def move_document(doc_id: str, tenant_id: str, folder_id: str | None) -> dict:
    """Move a document to a folder (or root if folder_id is None)."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "editor"):
        raise PermissionError("Access denied")
    CollaborationDocumentService.update_by_id(doc_id, {"folder_id": folder_id})
    return {"id": doc_id, "folder_id": folder_id}


# ── ACL Management ──

async def list_collaborators(doc_id: str, tenant_id: str) -> list[dict]:
    """List all collaborators for a document. Requires owner role."""
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")
    acls = CollaborationDocumentACLService.query(document_id=doc_id)
    return [
        {"id": a.id, "user_id": a.user_id, "role": a.role, "granted_by": a.granted_by, "create_time": a.create_time}
        for a in acls
    ]


async def add_collaborator(doc_id: str, tenant_id: str, user_id: str, role: str) -> dict:
    """Add a collaborator to a document. Requires owner role."""
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")
    if role not in ROLE_HIERARCHY:
        raise ValueError(f"Invalid role: {role}. Must be one of: {', '.join(ROLE_HIERARCHY.keys())}")
    acl_id = get_uuid()
    CollaborationDocumentACLService.save(
        id=acl_id,
        document_id=doc_id,
        user_id=user_id,
        role=role,
        granted_by=tenant_id,
        create_time=current_timestamp(),
    )
    return {"id": acl_id, "document_id": doc_id, "user_id": user_id, "role": role}


async def update_collaborator_role(doc_id: str, tenant_id: str, collaborator_user_id: str, role: str) -> dict:
    """Update a collaborator's role. Requires owner role."""
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")
    if role not in ROLE_HIERARCHY:
        raise ValueError(f"Invalid role: {role}. Must be one of: {', '.join(ROLE_HIERARCHY.keys())}")
    acls = CollaborationDocumentACLService.query(document_id=doc_id, user_id=collaborator_user_id)
    if not acls:
        raise LookupError("Collaborator not found")
    CollaborationDocumentACLService.update_by_id(acls[0].id, {"role": role})
    return {"document_id": doc_id, "user_id": collaborator_user_id, "role": role}


async def remove_collaborator(doc_id: str, tenant_id: str, collaborator_user_id: str) -> bool:
    """Remove a collaborator from a document. Requires owner role."""
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")
    acls = CollaborationDocumentACLService.query(document_id=doc_id, user_id=collaborator_user_id)
    if not acls:
        raise LookupError("Collaborator not found")
    CollaborationDocumentACLService.delete_by_id(acls[0].id)
    return True


# ── Comment CRUD ──

async def list_comments(doc_id: str, tenant_id: str) -> list[dict]:
    """List all non-deleted comments for a document."""
    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")

    comments = (
        CollaborationComment.select()
        .where(
            (CollaborationComment.document_id == doc_id)
            & (CollaborationComment.deleted_at.is_null())
        )
        .order_by(CollaborationComment.create_time.asc())
    )

    # Batch-fetch user names for all commenters
    user_ids = list({c.user_id for c in comments})
    user_map: dict[str, str] = {}
    if user_ids:
        users = User.select(User.id, User.nickname).where(User.id.in_(user_ids))
        user_map = {u.id: u.nickname for u in users}

    return [
        {
            "id": c.id,
            "document_id": c.document_id,
            "user_id": c.user_id,
            "user_name": user_map.get(c.user_id, c.user_id),
            "parent_comment_id": c.parent_comment_id,
            "anchor_block_key": c.anchor_block_key,
            "anchor_offset_start": c.anchor_offset_start,
            "anchor_offset_end": c.anchor_offset_end,
            "content": c.content,
            "resolved": c.resolved,
            "create_time": c.create_time,
            "update_time": c.update_time,
        }
        for c in comments
    ]


async def create_comment(
    doc_id: str,
    tenant_id: str,
    content: str,
    parent_comment_id: str | None = None,
    anchor_block_key: str | None = None,
    anchor_offset_start: int | None = None,
    anchor_offset_end: int | None = None,
) -> dict:
    """Create a comment on a document."""
    role = _get_user_role(doc_id, tenant_id)
    if not role or ROLE_HIERARCHY.get(role, 0) < ROLE_HIERARCHY.get("commenter", 1):
        raise PermissionError("Access denied")

    if not content or not content.strip():
        raise ValueError("Comment content is required")

    safe_content = sanitize_html(content.strip())

    comment_id = get_uuid()
    now = current_timestamp()
    CollaborationCommentService.save(
        id=comment_id,
        document_id=doc_id,
        user_id=tenant_id,
        parent_comment_id=parent_comment_id,
        anchor_block_key=anchor_block_key,
        anchor_offset_start=anchor_offset_start,
        anchor_offset_end=anchor_offset_end,
        content=safe_content,
        create_time=now,
        update_time=now,
    )
    # Fetch user name for the creator
    user_name = tenant_id
    try:
        users = User.select(User.id, User.nickname).where(User.id == tenant_id)
        if users:
            user_name = users[0].nickname
    except Exception:
        pass

    return {
        "id": comment_id,
        "document_id": doc_id,
        "user_id": tenant_id,
        "user_name": user_name,
        "parent_comment_id": parent_comment_id,
        "content": safe_content,
        "resolved": False,
        "create_time": now,
        "update_time": now,
    }


async def update_comment(doc_id: str, comment_id: str, tenant_id: str, content: str) -> dict:
    """Edit own comment content."""
    e, comment = CollaborationCommentService.get_by_id(comment_id)
    if not e:
        raise LookupError("Comment not found")
    if comment.document_id != doc_id:
        raise LookupError("Comment not found")
    if comment.user_id != tenant_id:
        raise PermissionError("Can only edit own comments")
    if comment.deleted_at is not None:
        raise LookupError("Comment has been deleted")
    if not content or not content.strip():
        raise ValueError("Comment content is required")

    safe_content = sanitize_html(content.strip())

    CollaborationCommentService.update_by_id(comment_id, {
        "content": safe_content,
        "update_time": current_timestamp(),
    })
    return {"id": comment_id, "content": safe_content}


async def delete_comment(doc_id: str, comment_id: str, tenant_id: str) -> bool:
    """Soft-delete own comment. Also soft-deletes all child comments."""
    e, comment = CollaborationCommentService.get_by_id(comment_id)
    if not e:
        raise LookupError("Comment not found")
    if comment.document_id != doc_id:
        raise LookupError("Comment not found")
    if comment.user_id != tenant_id:
        raise PermissionError("Can only delete own comments")

    now = current_timestamp()
    # Collect all descendant IDs iteratively, avoiding circular references
    ids_to_delete = {comment_id}
    child_ids = [comment_id]
    while child_ids:
        children = CollaborationComment.select().where(
            (CollaborationComment.parent_comment_id.in_(child_ids))
            & (CollaborationComment.deleted_at.is_null())
        )
        child_ids = []
        for c in children:
            if c.id not in ids_to_delete:
                child_ids.append(c.id)
                ids_to_delete.add(c.id)

    for cid in ids_to_delete:
        CollaborationCommentService.update_by_id(cid, {"deleted_at": now})
    return True


async def resolve_comment(doc_id: str, comment_id: str, tenant_id: str) -> dict:
    """Mark a comment as resolved."""
    e, comment = CollaborationCommentService.get_by_id(comment_id)
    if not e:
        raise LookupError("Comment not found")
    if comment.document_id != doc_id:
        raise LookupError("Comment not found")
    if comment.deleted_at is not None:
        raise LookupError("Comment has been deleted")

    CollaborationCommentService.update_by_id(comment_id, {"resolved": True})
    return {"id": comment_id, "resolved": True}


async def unresolve_comment(doc_id: str, comment_id: str, tenant_id: str) -> dict:
    """Reopen a resolved comment."""
    e, comment = CollaborationCommentService.get_by_id(comment_id)
    if not e:
        raise LookupError("Comment not found")
    if comment.document_id != doc_id:
        raise LookupError("Comment not found")
    if comment.deleted_at is not None:
        raise LookupError("Comment has been deleted")

    CollaborationCommentService.update_by_id(comment_id, {"resolved": False})
    return {"id": comment_id, "resolved": False}


# ── Version History ──

async def list_versions(doc_id: str, tenant_id: str) -> dict:
    """返回文档的版本历史。

    - current_version: 当前文档版本号
    - has_ydoc: 是否存在可恢复的 ydoc 状态
    - update_time: 主表最后一次更新时间
    - versions: 历史快照列表 (按时间倒序，最多 MAX_DOCUMENT_VERSIONS_KEEP 条)
    """
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")

    snapshots = (
        CollaborationDocumentVersion
        .select(
            CollaborationDocumentVersion.id,
            CollaborationDocumentVersion.version,
            CollaborationDocumentVersion.created_by,
            CollaborationDocumentVersion.create_time,
        )
        .where(CollaborationDocumentVersion.document_id == doc_id)
        .order_by(
            CollaborationDocumentVersion.create_time.desc(),
            CollaborationDocumentVersion.version.desc(),
        )
        .limit(MAX_DOCUMENT_VERSIONS_KEEP)
    )
    snap_list = list(snapshots)

    user_ids = list({s.created_by for s in snap_list if s.created_by})
    nickname_map = {}
    if user_ids:
        try:
            rows = User.select(User.id, User.nickname).where(User.id.in_(user_ids))
            nickname_map = {row.id: row.nickname for row in rows}
        except Exception:
            pass

    return {
        "current_version": doc.version or 0,
        "has_ydoc": bool(doc.ydoc),
        "update_time": doc.update_time,
        "versions": [
            {
                "id": s.id,
                "version": s.version,
                "created_by": s.created_by,
                "user_name": nickname_map.get(s.created_by, s.created_by),
                "create_time": s.create_time,
            }
            for s in snap_list
        ],
    }


async def restore_version(doc_id: str, tenant_id: str, version_id: str | None) -> dict:
    """恢复到指定历史版本。

    语义：恢复 ≠ 回滚。把目标快照的 ydoc + content 作为新版本写入文档
    (version 继续递增)，并写一条快照代表「恢复后的新当前状态」。
    反悔路径：上一次 save 的快照仍然在表里，用户点它即可回到恢复前。
    恢复后前端通过 reload 重新加载。
    """
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "editor"):
        raise PermissionError("Access denied")

    snapshots = (
        CollaborationDocumentVersion
        .select(
            CollaborationDocumentVersion.id,
            CollaborationDocumentVersion.version,
            CollaborationDocumentVersion.ydoc_snapshot,
            CollaborationDocumentVersion.content_snapshot,
        )
        .where(CollaborationDocumentVersion.document_id == doc_id)
        .order_by(
            CollaborationDocumentVersion.create_time.desc(),
            CollaborationDocumentVersion.version.desc(),
        )
        .limit(MAX_DOCUMENT_VERSIONS_KEEP)
    )
    snap_list = list(snapshots)

    if not snap_list:
        raise ValueError("No saved version to restore")

    target = None
    if version_id:
        target = next((s for s in snap_list if s.id == version_id), None)
        if target is None:
            # 严格匹配：前端指定的版本可能已被 trim 掉，静默回退到最新
            # 会让用户误以为恢复了目标版本。显式报错让前端 refresh 后重试。
            raise LookupError(f"Version {version_id} not found (may have been trimmed)")
    else:
        # 未指定 → 取最新一条 (旧行为兼容)
        target = snap_list[0]

    now_ms = current_timestamp()
    new_version = (doc.version or 0) + 1

    target_ydoc = target.ydoc_snapshot
    target_content = target.content_snapshot

    # 1) 把目标快照写入主表。target 的 ydoc/content 始终来自历史 save，
    #    save 时刻保证 ydoc 非空 (save_ydoc_state 只在 ydoc_state 存在时写快照)。
    update_data = {
        "ydoc": target_ydoc,
        "content": target_content,
        "version": new_version,
    }
    CollaborationDocumentService.update_by_id(doc_id, update_data)

    # 2) 写一条快照代表「恢复后的当前状态」，与主表 version 对齐，
    #    前端 list_versions 据此显示"当前版本"标记。
    try:
        CollaborationDocumentVersionService.save(
            id=get_uuid(),
            document_id=doc_id,
            version=new_version,
            ydoc_snapshot=target_ydoc,
            content_snapshot=target_content,
            created_by=tenant_id,
            create_time=now_ms,
        )
        _trim_document_versions(doc_id)
    except Exception as exc:
        logging.warning(f"[restore_version] post-restore snapshot failed for doc {doc_id}: {exc}")

    await log_audit(tenant_id, doc_id, "version.restore", {
        "version": new_version,
        "restored_from_version": target.version,
    })

    return {
        "id": doc_id,
        "version": new_version,
        "restored_from_version": target.version,
        "ydoc": base64.b64encode(target_ydoc).decode("ascii") if target_ydoc else None,
        "content": target_content,
    }


# ── Share Link ──

async def create_or_update_share(doc_id: str, tenant_id: str, permission: str = "view", password: str = None, expires_at: int = None) -> dict:
    """Create or update a share link for a document. One share per document."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")

    if permission not in ("view", "edit"):
        raise ValueError("Permission must be 'view' or 'edit'")

    password_hash = generate_password_hash(password) if password else None

    # Check if share already exists
    existing = CollaborationShareLinkService.query(document_id=doc_id)
    if existing:
        share = existing[0]
        update_data = {"permission": permission, "update_time": current_timestamp()}
        if password is not None:
            # password="" means clear password, password="xxx" means set new password
            update_data["password_hash"] = generate_password_hash(password) if password else None
        if expires_at is not None:
            update_data["expires_at"] = expires_at
        CollaborationShareLinkService.update_by_id(share.id, update_data)
        token = share.token
        await log_audit(tenant_id, doc_id, "share.update", {"permission": permission})
    else:
        token = get_uuid()
        CollaborationShareLinkService.save(
            id=get_uuid(),
            document_id=doc_id,
            token=token,
            permission=permission,
            password_hash=generate_password_hash(password) if password else None,
            expires_at=expires_at,
            created_by=tenant_id,
            create_time=current_timestamp(),
            update_time=current_timestamp(),
        )
        await log_audit(tenant_id, doc_id, "share.create", {"permission": permission})

    return {
        "document_id": doc_id,
        "token": token,
        "permission": permission,
        "has_password": bool(password),
        "expires_at": expires_at,
    }


async def get_share(doc_id: str, tenant_id: str) -> dict | None:
    """Get current share link info for a document."""
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")

    existing = CollaborationShareLinkService.query(document_id=doc_id)
    if not existing:
        return None

    share = existing[0]
    return {
        "document_id": share.document_id,
        "token": share.token,
        "permission": share.permission,
        "has_password": bool(share.password_hash),
        "expires_at": share.expires_at,
    }


async def delete_share(doc_id: str, tenant_id: str) -> bool:
    """Delete the share link for a document."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")

    existing = CollaborationShareLinkService.query(document_id=doc_id)
    if not existing:
        raise LookupError("Share link not found")

    CollaborationShareLinkService.delete_by_id(existing[0].id)
    await log_audit(tenant_id, doc_id, "share.delete")

    return True


async def access_shared_doc(token: str, password: str = None) -> dict:
    """Access a shared document by token. Returns document data if access is granted."""
    existing = CollaborationShareLinkService.query(token=token)
    if not existing:
        raise LookupError("Share link not found or expired")

    share = existing[0]

    # Check expiry
    if share.expires_at and share.expires_at < current_timestamp():
        raise PermissionError("Share link has expired")

    # Check password
    if share.password_hash:
        if not password:
            raise PermissionError("Password required")
        if not check_password_hash(share.password_hash, password):
            raise PermissionError("Incorrect password")

    e, doc = CollaborationDocumentService.get_by_id(share.document_id)
    if not e:
        raise LookupError("Document not found")

    return {
        "id": doc.id,
        "name": doc.name,
        "content": doc.content,
        "markdown_content": doc.markdown_content,
        "permission": share.permission,
        "has_password": bool(share.password_hash),
    }


# ── Attachment CRUD ──

async def upload_attachment(doc_id: str, tenant_id: str, file_obj) -> dict:
    """Upload an attachment to a document. Stored in MinIO via STORAGE_IMPL."""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _check_role(doc_id, tenant_id, "editor"):
        raise PermissionError("Access denied")

    data, safe_name, mime_type = validate_upload(file_obj)

    attachment_id = get_uuid()
    ext = f".{safe_name.rsplit('.', 1)[-1].lower()}" if "." in safe_name else ""
    storage_key = f"attachments/{doc_id}/{attachment_id}{ext}"

    try:
        settings.STORAGE_IMPL.put("collaboration", storage_key, data)
    except Exception as e:
        logging.error(f"Failed to upload attachment to storage: {e}")
        raise RuntimeError("Failed to store attachment")

    CollaborationAttachmentService.save(
        id=attachment_id,
        document_id=doc_id,
        file_name=safe_name,
        file_size=len(data),
        mime_type=mime_type,
        storage_key=storage_key,
        uploader_id=tenant_id,
        create_time=current_timestamp(),
    )

    await log_audit(tenant_id, doc_id, "attachment.upload", {"file_name": safe_name, "file_size": len(data)})

    return {
        "id": attachment_id,
        "document_id": doc_id,
        "file_name": safe_name,
        "file_size": len(data),
        "mime_type": mime_type,
    }


async def list_attachments(doc_id: str, tenant_id: str) -> list[dict]:
    """List all attachments for a document."""
    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")

    attachments = CollaborationAttachmentService.query(document_id=doc_id)
    return [
        {
            "id": a.id,
            "document_id": a.document_id,
            "file_name": a.file_name,
            "file_size": a.file_size,
            "mime_type": a.mime_type,
            "uploader_id": a.uploader_id,
            "create_time": a.create_time,
        }
        for a in attachments
    ]


async def download_attachment(doc_id: str, attachment_id: str, tenant_id: str) -> tuple[bytes, str, str]:
    """Download an attachment. Returns (data, filename, mime_type)."""
    e, attachment = CollaborationAttachmentService.get_by_id(attachment_id)
    if not e:
        raise LookupError("Attachment not found")
    if attachment.document_id != doc_id:
        raise LookupError("Attachment not found")

    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")

    try:
        data = settings.STORAGE_IMPL.get("collaboration", attachment.storage_key)
    except Exception as e:
        logging.error(f"Failed to download attachment: {e}")
        raise RuntimeError("Failed to retrieve attachment")

    return data, attachment.file_name, attachment.mime_type


# ── Spreadsheet image asset proxy (P4) ─────────────────────────────────
# xlsx adapter uploads floating images to MinIO bucket 'collaboration' under
# key 'docs/{doc_id}/images/{asset_id}.{ext}'. The frontend Univer drawing
# plugin renders them via <img src> which can't carry Authorization headers,
# so the asset endpoint authenticates via ?token=<jwt> query param.
_EXT_TO_MIME = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "webp": "image/webp",
    "svg": "image/svg+xml",
    "tiff": "image/tiff",
    "ico": "image/x-icon",
}


def list_doc_asset_ids(doc_id: str) -> list[str]:
    """Return all asset_ids (uuid) for a doc, by scanning its storage prefix.

    Used by document delete to clean up orphaned images in MinIO.
    """
    # The STORAGE_IMPL interface has no list() method; we rely on callers
    # tracking assets via the workbook's SHEET_DRAWING_PLUGIN resource.
    return []


def get_doc_asset(doc_id: str, asset_id: str, tenant_id: str) -> tuple[bytes, str]:
    """Fetch a spreadsheet image asset from MinIO.

    Returns (bytes, mimetype). Raises LookupError if not found, PermissionError
    if the caller lacks read access to the document.
    """
    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")

    # Try each supported extension — the asset_id alone doesn't carry the ext.
    for ext, mime in _EXT_TO_MIME.items():
        storage_key = f"docs/{doc_id}/images/{asset_id}.{ext}"
        try:
            data = settings.STORAGE_IMPL.get("collaboration", storage_key)
            if data:
                return data, mime
        except Exception:
            continue
    raise LookupError(f"Asset {asset_id} not found for document {doc_id}")


async def delete_attachment(doc_id: str, attachment_id: str, tenant_id: str) -> bool:
    """Delete an attachment."""
    e, attachment = CollaborationAttachmentService.get_by_id(attachment_id)
    if not e:
        raise LookupError("Attachment not found")
    if attachment.document_id != doc_id:
        raise LookupError("Attachment not found")
    if not _check_role(doc_id, tenant_id, "editor"):
        raise PermissionError("Access denied")

    # Delete DB record first (source of truth), then best-effort storage cleanup
    CollaborationAttachmentService.delete_by_id(attachment_id)

    try:
        settings.STORAGE_IMPL.rm("collaboration", attachment.storage_key)
    except Exception as e:
        logging.warning(f"Failed to delete attachment from storage (orphaned): {e}")

    await log_audit(tenant_id, doc_id, "attachment.delete", {"file_name": attachment.file_name})

    return True


# ── Audit Log ──

async def list_audit_logs(doc_id: str, tenant_id: str, limit: int = 50, offset: int = 0) -> dict:
    """List audit logs for a document. Owner only."""
    if not _check_role(doc_id, tenant_id, "owner"):
        raise PermissionError("Access denied")

    logs = (
        CollaborationAuditLog.select(
            CollaborationAuditLog.id,
            CollaborationAuditLog.user_id,
            CollaborationAuditLog.action,
            CollaborationAuditLog.detail,
            CollaborationAuditLog.ip_address,
            CollaborationAuditLog.create_time,
        )
        .where(CollaborationAuditLog.document_id == doc_id)
        .order_by(CollaborationAuditLog.create_time.desc())
        .offset(offset)
        .limit(limit)
    )
    total = (
        CollaborationAuditLog.select()
        .where(CollaborationAuditLog.document_id == doc_id)
        .count()
    )

    # Collect unique user_ids and resolve nicknames
    user_ids = list({log.user_id for log in logs if log.user_id})
    nickname_map = {}
    if user_ids:
        rows = (
            User.select(UserTenant.user_id, User.nickname)
            .join(UserTenant, on=(User.id == UserTenant.user_id))
            .where(UserTenant.tenant_id.in_(user_ids))
        )
        nickname_map = {row.user_id: row.nickname for row in rows}

    return {
        "total": total,
        "logs": [
            {
                "id": log.id,
                "user_id": log.user_id,
                "user_name": nickname_map.get(log.user_id, log.user_id),
                "action": log.action,
                "detail": log.detail,
                "ip_address": log.ip_address,
                "create_time": log.create_time,
            }
            for log in logs
        ],
    }


async def save_exported_file(doc_id: str, tenant_id: str, blob: bytes, fmt: str) -> dict:
    """前端导出 docx/pdf 后上传 blob，存到 STORAGE_IMPL，更新 file_path。

    后端不做任何格式生成，只存文件。Univer Docs 的导出在前端浏览器跑。
    """
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")
    if fmt not in ("docx", "pdf"):
        raise ValueError(f"Unsupported format: {fmt}")
    storage_key = f"{doc_id}.{fmt}"
    settings.STORAGE_IMPL.put("collaboration", storage_key, blob)
    CollaborationDocumentService.update_by_id(
        doc_id, {"file_path": storage_key, "file_type": fmt}
    )
    return {"file_path": storage_key, "size": len(blob)}


async def get_exported_file(doc_id: str, tenant_id: str) -> tuple | None:
    """返回最近一次导出的 (blob_bytes, filename, mimetype)，无则 None。"""
    e, doc = CollaborationDocumentService.get_by_id(doc_id)
    if not e:
        raise LookupError("Document not found")
    if not _get_user_role(doc_id, tenant_id):
        raise PermissionError("Access denied")
    storage_key = doc.file_path
    if not storage_key:
        return None
    ext = storage_key.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif ext == "pdf":
        mimetype = "application/pdf"
    else:
        mimetype = "application/octet-stream"
    blob = settings.STORAGE_IMPL.get("collaboration", storage_key)
    filename = f"{doc.name}.{ext}"
    return blob, filename, mimetype
